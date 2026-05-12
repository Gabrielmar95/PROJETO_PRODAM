"""
detran_pipeline_nao_pdf.py — Processa arquivos não-PDF de DETRAN_CONSOLIDADO.

Tipos suportados:
  .docx  → python-docx (texto + tabelas)
  .xlsx  → openpyxl (cada aba como dict)
  .xls   → xlrd
  .html  → BeautifulSoup (texto limpo)
  .md    → texto direto
  .py    → texto direto + metadata (linhas, funções)
  .png/.jpg → pytesseract OCR
  .json  → copia conteúdo
  Thumbs.db, .DS_Store → ignorados

Saída: DETRAN_CONSOLIDADO_JSON/<categoria>/<filename>.json
"""
from __future__ import annotations
import sys, json, re
from pathlib import Path
from datetime import date
from collections import defaultdict

sys.stdout.reconfigure(encoding='utf-8')
ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "scripts"))
from prodam_utils import norm
from logging_config import get_logger
_logger = get_logger(__name__)

SRC = ROOT / "PRODAM_DOCS" / "DETRAN_CONSOLIDADO"
OUT = ROOT / "DETRAN_CONSOLIDADO_JSON"

import docx
import openpyxl
import xlrd
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image

RE_CNPJ = re.compile(r"\b(\d{2}[.\s]\d{3}[.\s]\d{3}[/\s]\d{4}[-\s]\d{2})\b")
RE_VALOR = re.compile(r"R\$\s*([\d.]+(?:[,.]\d{2})?)")
RE_DATA = re.compile(r"\b(\d{1,2}/\d{1,2}/\d{2,4})\b")
RE_NE = re.compile(r"\b(\d{4}NE\d{4,7})\b", re.IGNORECASE)
RE_INDICES = {k: re.compile(rf"\b{k}\b", re.IGNORECASE) for k in ("IGPM","IGP-M","IPCA","INPC","SELIC")}

def parse_comum(texto: str) -> dict:
    out = {}
    if cnpjs := list(dict.fromkeys(RE_CNPJ.findall(texto))):
        out["cnpjs"] = cnpjs[:5]
    if datas := list(dict.fromkeys(RE_DATA.findall(texto))):
        out["datas_encontradas"] = datas[:10]
    if valores := RE_VALOR.findall(texto):
        out["valores_brl"] = list(dict.fromkeys(valores))[:15]
    if nes := list(dict.fromkeys(m.group(1) for m in RE_NE.finditer(texto))):
        out["nes_referenciadas"] = nes[:15]
    indices = [k.replace("-","") for k,rx in RE_INDICES.items() if rx.search(texto)]
    if indices: out["indices_mencionados"] = list(dict.fromkeys(indices))
    return out

# ============================================================
# HANDLERS POR TIPO
# ============================================================
def handle_docx(path: Path) -> dict:
    d = docx.Document(str(path))
    paragrafos = [p.text for p in d.paragraphs if p.text.strip()]
    texto = "\n".join(paragrafos)
    tabelas = []
    for t in d.tables:
        rows = []
        for row in t.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tabelas.append(rows)
    return {
        "metodo": "python-docx",
        "num_paragrafos": len(paragrafos),
        "num_tabelas": len(tabelas),
        "texto_completo": texto,
        "tabelas": tabelas[:5],
        "texto_length": len(texto),
    }

def handle_xlsx(path: Path) -> dict:
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    abas = {}
    texto_agregado = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None and str(c).strip() for c in row):
                rows.append([str(c) if c is not None else "" for c in row])
        abas[sheet_name] = {
            "total_linhas": len(rows),
            "num_colunas": max((len(r) for r in rows), default=0),
            "amostra": rows[:10],
            "dados_completos": rows,
        }
        texto_agregado.append(f"=== ABA: {sheet_name} ===")
        for r in rows:
            texto_agregado.append(" | ".join(r))
    wb.close()
    texto = "\n".join(texto_agregado)
    return {
        "metodo": "openpyxl",
        "num_abas": len(abas),
        "abas": abas,
        "texto_completo": texto,
        "texto_length": len(texto),
    }

def handle_xls(path: Path) -> dict:
    wb = xlrd.open_workbook(str(path))
    abas = {}
    texto_agregado = []
    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        rows = []
        for i in range(ws.nrows):
            row = [str(ws.cell_value(i, j)) for j in range(ws.ncols)]
            if any(c.strip() for c in row):
                rows.append(row)
        abas[sheet_name] = {"total_linhas": len(rows), "amostra": rows[:10], "dados_completos": rows}
        texto_agregado.append(f"=== ABA: {sheet_name} ===")
        for r in rows: texto_agregado.append(" | ".join(r))
    texto = "\n".join(texto_agregado)
    return {
        "metodo": "xlrd",
        "num_abas": len(abas),
        "abas": abas,
        "texto_completo": texto,
        "texto_length": len(texto),
    }

def handle_html(path: Path) -> dict:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    # Remove scripts/styles
    for tag in soup(["script", "style", "meta", "link"]):
        tag.decompose()
    texto = soup.get_text(separator="\n", strip=True)
    titulo = soup.title.string if soup.title else None
    # Tabelas HTML
    tabelas_html = []
    for table in soup.find_all("table"):
        rows = []
        for tr in table.find_all("tr"):
            rows.append([td.get_text(strip=True) for td in tr.find_all(["td","th"])])
        tabelas_html.append(rows)
    return {
        "metodo": "beautifulsoup",
        "titulo": titulo,
        "num_tabelas": len(tabelas_html),
        "tabelas": tabelas_html[:5],
        "texto_completo": texto,
        "texto_length": len(texto),
    }

def handle_md(path: Path) -> dict:
    texto = path.read_text(encoding="utf-8", errors="ignore")
    # Conta headers
    headers = re.findall(r"^(#{1,6})\s+(.+)$", texto, re.MULTILINE)
    num_linhas = texto.count("\n")
    return {
        "metodo": "texto_direto",
        "num_linhas": num_linhas,
        "num_headers": len(headers),
        "titulos": [h[1][:100] for h in headers[:20]],
        "texto_completo": texto,
        "texto_length": len(texto),
    }

def handle_py(path: Path) -> dict:
    texto = path.read_text(encoding="utf-8", errors="ignore")
    funcoes = re.findall(r"^def\s+(\w+)\s*\(", texto, re.MULTILINE)
    classes = re.findall(r"^class\s+(\w+)", texto, re.MULTILINE)
    imports = re.findall(r"^(?:import|from)\s+(\S+)", texto, re.MULTILINE)
    return {
        "metodo": "texto_direto",
        "num_linhas": texto.count("\n"),
        "num_funcoes": len(funcoes),
        "funcoes": funcoes[:30],
        "num_classes": len(classes),
        "classes": classes[:10],
        "imports_top10": list(dict.fromkeys(imports))[:10],
        "texto_completo": texto,
        "texto_length": len(texto),
    }

def handle_image(path: Path) -> dict:
    try:
        img = Image.open(str(path))
        texto = pytesseract.image_to_string(img, config="--psm 3 -l por")
        return {
            "metodo": "ocr_tesseract_png",
            "tamanho": f"{img.width}x{img.height}",
            "texto_completo": texto,
            "texto_length": len(texto),
        }
    except Exception as e:
        return {"metodo": "imagem", "erro": str(e), "texto_completo": "", "texto_length": 0}

def handle_json(path: Path) -> dict:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return {
            "metodo": "json_direto",
            "tipo_raiz": type(data).__name__,
            "num_chaves" if isinstance(data, dict) else "num_itens": len(data) if hasattr(data,"__len__") else None,
            "dados": data,
            "texto_completo": json.dumps(data, ensure_ascii=False, indent=2)[:10000],
            "texto_length": len(json.dumps(data, ensure_ascii=False)),
        }
    except Exception as e:
        return {"metodo": "json", "erro": str(e)}

# ============================================================
# ROUTER
# ============================================================
HANDLERS = {
    ".docx": handle_docx,
    ".xlsx": handle_xlsx,
    ".xls": handle_xls,
    ".html": handle_html,
    ".htm": handle_html,
    ".md": handle_md,
    ".py": handle_py,
    ".png": handle_image,
    ".jpg": handle_image,
    ".jpeg": handle_image,
    ".json": handle_json,
}
IGNORAR = {".db", ".ini", ".DS_Store", ".lnk", ".tmp"}
IGNORAR_NAMES = {"Thumbs.db", ".DS_Store", "desktop.ini"}

def safe_filename(name):
    return re.sub(r"[^\w.-]", "_", name)

# ============================================================
# MAIN
# ============================================================
def processar_arquivo(path: Path, categoria: str) -> dict:
    ext = path.suffix.lower()
    if ext in IGNORAR or path.name in IGNORAR_NAMES:
        return None
    if ext == ".pdf":
        return None  # já tratado por pipeline_fast
    if ext == ".csv":
        return None  # já tratado

    handler = HANDLERS.get(ext)
    if not handler:
        return {"filename": path.name, "erro": f"extensão não suportada: {ext}", "tipo_doc": "UNKNOWN"}

    try:
        resultado = handler(path)
        texto = resultado.get("texto_completo", "")
        campos = parse_comum(texto) if texto else {}

        # Tipo doc baseado na extensão e categoria
        tipo_map = {
            ".docx": "DOCX", ".xlsx": "PLANILHA", ".xls": "PLANILHA",
            ".html": "HTML_RELATORIO", ".md": "MARKDOWN", ".py": "SCRIPT_PYTHON",
            ".png": "IMAGEM_OCR", ".json": "JSON_DADOS",
        }
        return {
            "filename": path.name,
            "path": str(path.relative_to(ROOT)),
            "categoria": categoria,
            "extensao": ext,
            "data_extracao": date.today().isoformat(),
            "tamanho_bytes": path.stat().st_size,
            "tipo_doc": tipo_map.get(ext, "OUTRO"),
            "texto_preview": texto[:800],
            "campos_estruturados": campos,
            **resultado,
        }
    except Exception as e:
        return {"filename": path.name, "categoria": categoria, "erro": str(e)[:200]}

# ============================================================
# Coleta e processa
# ============================================================
tarefas = []
for cat_folder in ("06_COBRANCAS", "09_RELATORIOS"):
    folder = SRC / cat_folder
    if not folder.exists(): continue
    out_folder = OUT / cat_folder
    out_folder.mkdir(exist_ok=True)
    for f in folder.iterdir():
        if not f.is_file(): continue
        if f.suffix.lower() == ".pdf": continue
        if f.suffix.lower() == ".csv": continue
        if f.name in IGNORAR_NAMES: continue
        if f.suffix.lower() in IGNORAR: continue
        json_out = out_folder / (safe_filename(f.stem) + f.suffix.replace(".","_") + ".json")
        # Evitar duplicar JSONs já feitos
        if json_out.exists() and json_out.stat().st_size > 500:
            continue
        tarefas.append((f, cat_folder, json_out))

print(f"Arquivos não-PDF a processar: {len(tarefas)}\n")

resultados = []
for i, (f, cat, json_out) in enumerate(tarefas, 1):
    r = processar_arquivo(f, cat)
    if r is None:
        continue
    resultados.append(r)
    try:
        json_out.write_text(json.dumps(r, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
        status = "OK" if not r.get("erro") else "ERR"
        print(f"  [{i:2d}/{len(tarefas)}] [{status}] {f.name[:55]:<55} → {r.get('metodo','?')}")
    except Exception as e:
        print(f"  [{i:2d}/{len(tarefas)}] [SAVE_ERR] {f.name}: {e}")

# ============================================================
# Atualiza MASTER
# ============================================================
master_path = OUT / "MASTER_DETRAN_CONSOLIDADO.json"
master = json.loads(master_path.read_text(encoding="utf-8"))

all_jsons = []
for folder in OUT.iterdir():
    if folder.is_dir():
        for jf in folder.glob("*.json"):
            try: all_jsons.append(json.loads(jf.read_text(encoding="utf-8")))
            except Exception as exc: _logger.warning("JSON corrompido %s: %s", jf.name, exc)

master["total_arquivos"] = len(all_jsons)
master["data_atualizacao"] = date.today().isoformat()
por_cat = defaultdict(int)
por_ext = defaultdict(int)
for d in all_jsons:
    por_cat[d.get("categoria","?")] += 1
    ext = d.get("extensao") or Path(d.get("filename","")).suffix.lower()
    por_ext[ext] += 1

master["por_categoria"] = dict(por_cat)
master["por_extensao"] = dict(por_ext)

json.dump(master, open(master_path, "w", encoding="utf-8"), ensure_ascii=False, indent=2, default=str)

print(f"\n[MASTER] {master['total_arquivos']} JSONs indexados")
print(f"  Por categoria: {dict(por_cat)}")
print(f"  Por extensão:  {dict(por_ext)}")
