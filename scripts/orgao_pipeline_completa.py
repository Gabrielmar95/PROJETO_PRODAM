"""
orgao_pipeline_completa.py — Pipeline genérica para qualquer órgão do PRODAM.

Replica o que fizemos com DETRAN, parametrizado por --orgao SIGLA.
Aplicável aos 78 devedores com pasta *_CONSOLIDADO em PRODAM_DOCS/.

FASES:
  1. PDF→JSON fast-path (PyMuPDF)
  2. OCR fase 2 (Tesseract — PDFs imagem)
  3. Não-PDF (docx, xlsx, html, md, py, png)
  4. Ingestão de contratos no DB (PDF→spcf_contratos)
  5. Análise de índices de correção por contrato
  6. Score composto (12 dimensões)
  7. Relatórios consolidados

Uso:
  py -3.12 orgao_pipeline_completa.py --orgao SEDUC
  py -3.12 orgao_pipeline_completa.py --orgao "SES/SUSAM"
  py -3.12 orgao_pipeline_completa.py --orgao DETRAN --pular-ocr  # se OCR já rodado
  py -3.12 orgao_pipeline_completa.py --orgao SSP --so-fase 1,3    # só fases 1 e 3
  py -3.12 orgao_pipeline_completa.py --listar                     # lista órgãos disponíveis
"""
from __future__ import annotations
import sys, json, re, csv, time, sqlite3, argparse
from pathlib import Path
from datetime import date
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import defaultdict, Counter

sys.stdout.reconfigure(encoding='utf-8')
ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "scripts"))
from prodam_utils import norm, norm_variants, fmt_brl, parse_br_date, parse_comp, esta_prescrita

DOCS = ROOT / "PRODAM_DOCS"
DB_PATH = ROOT / "prodam.db"

# Handlers - importa com fallback
try:
    import fitz
except ImportError:
    fitz = None
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_OK = True
except ImportError:
    OCR_OK = False
try:
    import docx as pydocx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
try:
    import openpyxl
    XLSX_OK = True
except ImportError:
    XLSX_OK = False
try:
    import xlrd
    XLS_OK = True
except ImportError:
    XLS_OK = False
try:
    from bs4 import BeautifulSoup
    BS4_OK = True
except ImportError:
    BS4_OK = False

TESS_CONFIG = "--psm 3 -l por"

# Regex comuns
RE_VALOR = re.compile(r"R\$\s*([\d.]+(?:[,.]\d{2})?)")
RE_DATA = re.compile(r"\b(\d{1,2}/\d{1,2}/\d{2,4})\b")
RE_CNPJ = re.compile(r"\b(\d{2}[.\s]\d{3}[.\s]\d{3}[/\s]\d{4}[-\s]\d{2})\b")
RE_NE = re.compile(r"\b(\d{4}NE\d{4,7})\b", re.IGNORECASE)
RE_INDICES = {k: re.compile(rf"\b{k}\b", re.IGNORECASE) for k in ("IGPM","IGP-M","IPCA","INPC","SELIC")}
RE_CONTRATO_FN = re.compile(r"\b(\d{2,3})[.\s](\d{2,4})\b")
RE_OFICIO = re.compile(r"\b(?:Of[íi]cio|OF\.?)\s*(?:Nº?|n\.?)?\s*(\d+[/\-.]?\d{2,4})", re.IGNORECASE)

def parse_comum(texto: str) -> dict:
    out = {}
    if cnpjs := list(dict.fromkeys(RE_CNPJ.findall(texto))): out["cnpjs"] = cnpjs[:5]
    if datas := list(dict.fromkeys(RE_DATA.findall(texto))): out["datas_encontradas"] = datas[:10]
    if valores := RE_VALOR.findall(texto): out["valores_brl"] = list(dict.fromkeys(valores))[:15]
    if nes := list(dict.fromkeys(m.group(1) for m in RE_NE.finditer(texto))): out["nes_referenciadas"] = nes[:15]
    idxs = [k.replace("-","") for k,rx in RE_INDICES.items() if rx.search(texto)]
    if idxs: out["indices_mencionados"] = list(dict.fromkeys(idxs))
    return out

def safe_filename(name):
    return re.sub(r"[^\w.-]", "_", name)

def sigla_to_folder(sigla: str) -> str:
    return re.sub(r"[^\w]", "_", sigla).strip("_")

# ============================================================
# DETECTA PASTA DO ÓRGÃO
# ============================================================
def localizar_pasta_consolidado(orgao: str) -> Path | None:
    safe = sigla_to_folder(orgao)
    candidatos = [
        DOCS / f"{safe}_CONSOLIDADO",
        DOCS / f"{orgao.upper()}_CONSOLIDADO",
        DOCS / f"{orgao.replace('/','_').upper()}_CONSOLIDADO",
    ]
    for c in candidatos:
        if c.exists() and c.is_dir(): return c
    # Busca fuzzy
    alvo_norm = norm(orgao)
    for p in DOCS.iterdir():
        if p.is_dir() and p.name.endswith("_CONSOLIDADO"):
            nome_base = p.name[:-len("_CONSOLIDADO")]
            if norm(nome_base) == alvo_norm:
                return p
    return None

def listar_orgaos() -> list:
    return sorted([p.name[:-len("_CONSOLIDADO")] for p in DOCS.iterdir()
                   if p.is_dir() and p.name.endswith("_CONSOLIDADO")])

# ============================================================
# PARSERS POR CATEGORIA
# ============================================================
def parse_contrato(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "TERMO_ADITIVO" if re.search(r"\d[°ºo]\s*TA|aditivo", filename, re.IGNORECASE) else "CONTRATO_PRINCIPAL"
    m = re.search(r"cl[áa]usula\s+(?:d[ée]cima\s+primeira|11[°ªa]?|onze).{0,600}", texto, re.IGNORECASE | re.DOTALL)
    if m: out["clausula_11"] = m.group(0)[:500]
    m = re.search(r"cl[áa]usula\s+primeira.{0,400}", texto, re.IGNORECASE | re.DOTALL)
    if m: out["clausula_objeto"] = m.group(0)[:350]
    pcts = re.findall(r"(?:reajuste|atualiza).{0,80}?(\d{1,2}[,.]?\d{1,3})\s*%", texto, re.IGNORECASE | re.DOTALL)
    if pcts: out["reajustes_pct"] = [float(p.replace(",",".")) for p in pcts[:5]]
    return out

def parse_empenho(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "NOTA_EMPENHO"
    m = re.search(r"(\d{4}NE\d+)", texto, re.IGNORECASE)
    if m: out["numero_ne"] = m.group(1)
    m = re.search(r"Descri[çc][ãa]o[:\s]+(.+?)(?:\n\s*\n|Situa|Evento)", texto, re.IGNORECASE | re.DOTALL)
    if m: out["descricao"] = m.group(1).strip()[:400]
    return out

def parse_fatura(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "FATURA"
    m = re.search(r"Compet[êe]ncia[:\s]+(\d{1,2}/\d{4})", texto, re.IGNORECASE)
    if m: out["competencia"] = m.group(1)
    m = re.search(r"NFS?[e\-]?\s*(?:Nº|n\.?)?\s*(\d{4,7})", texto, re.IGNORECASE)
    if m: out["numero_nfse"] = m.group(1)
    return out

def parse_aceite(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "ACEITE_TECNICO"
    return out

def parse_cobranca(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "OFICIO_COBRANCA"
    m = RE_OFICIO.search(texto)
    if m: out["oficio_numero"] = m.group(1)
    return out

def parse_certidao(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "CERTIDAO"
    if "negativa" in texto.lower(): out["tipo_cnd"] = "NEGATIVA"
    elif "positiva" in texto.lower(): out["tipo_cnd"] = "POSITIVA"
    return out

def parse_rel(texto, filename):
    out = parse_comum(texto)
    out["tipo_doc"] = "RELATORIO"
    return out

PARSERS = {
    "01_CONTRATOS": parse_contrato,
    "02_EMPENHOS": parse_empenho,
    "03_FATURAS": parse_fatura,
    "05_ACEITES": parse_aceite,
    "06_COBRANCAS": parse_cobranca,
    "08_PDFS_ORIGINAIS": parse_certidao,
    "09_RELATORIOS": parse_rel,
}

# ============================================================
# HANDLERS DE ARQUIVO
# ============================================================
def extrair_pymupdf(path):
    meta = {"metodo": "pymupdf", "paginas": 0, "tem_texto": False}
    try:
        with fitz.open(str(path)) as doc:
            meta["paginas"] = len(doc)
            textos = [pg.get_text() for pg in doc]
            texto = "\n".join(textos)
            meta["tem_texto"] = len(texto.strip()) > 100
        return texto, meta
    except Exception as e:
        meta["erro"] = str(e)[:150]
        return "", meta

def ocr_pdf(path):
    meta = {"metodo": "ocr_tesseract", "paginas_ocr": 0}
    try:
        images = convert_from_path(str(path), dpi=200, grayscale=True)
        textos = [pytesseract.image_to_string(img, config=TESS_CONFIG) for img in images]
        meta["paginas_ocr"] = len(images)
        return "\n\n===PAGINA===\n\n".join(textos), meta
    except Exception as e:
        meta["erro"] = str(e)[:150]
        return "", meta

def handle_docx(path):
    d = pydocx.Document(str(path))
    paragrafos = [p.text for p in d.paragraphs if p.text.strip()]
    texto = "\n".join(paragrafos)
    tabelas = [[[c.text.strip() for c in r.cells] for r in t.rows] for t in d.tables]
    return texto, {"metodo":"python-docx","num_paragrafos":len(paragrafos),"num_tabelas":len(tabelas),"tabelas":tabelas[:5]}

def handle_xlsx(path):
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    abas = {}; texto = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = [[str(c) if c is not None else "" for c in r] for r in ws.iter_rows(values_only=True) if any(c is not None and str(c).strip() for c in r)]
        abas[sn] = {"total_linhas": len(rows), "amostra": rows[:10], "dados_completos": rows}
        texto.append(f"=== {sn} ==="); texto.extend(" | ".join(r) for r in rows)
    wb.close()
    return "\n".join(texto), {"metodo":"openpyxl","num_abas":len(abas),"abas":abas}

def handle_xls(path):
    wb = xlrd.open_workbook(str(path))
    abas = {}; texto = []
    for sn in wb.sheet_names():
        ws = wb.sheet_by_name(sn)
        rows = []
        for i in range(ws.nrows):
            row = [str(ws.cell_value(i,j)) for j in range(ws.ncols)]
            if any(c.strip() for c in row): rows.append(row)
        abas[sn] = {"total_linhas": len(rows), "amostra": rows[:10], "dados_completos": rows}
        texto.append(f"=== {sn} ==="); texto.extend(" | ".join(r) for r in rows)
    return "\n".join(texto), {"metodo":"xlrd","num_abas":len(abas),"abas":abas}

def handle_html(path):
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script","style","meta","link"]): tag.decompose()
    texto = soup.get_text(separator="\n", strip=True)
    tabelas = [[[td.get_text(strip=True) for td in tr.find_all(["td","th"])] for tr in table.find_all("tr")] for table in soup.find_all("table")]
    return texto, {"metodo":"beautifulsoup","num_tabelas":len(tabelas),"tabelas":tabelas[:5]}

def handle_md(path):
    texto = path.read_text(encoding="utf-8", errors="ignore")
    headers = re.findall(r"^(#{1,6})\s+(.+)$", texto, re.MULTILINE)
    return texto, {"metodo":"texto_direto","num_headers":len(headers),"titulos":[h[1][:100] for h in headers[:20]]}

def handle_py(path):
    texto = path.read_text(encoding="utf-8", errors="ignore")
    return texto, {
        "metodo":"texto_direto",
        "num_funcoes": len(re.findall(r"^def\s+(\w+)", texto, re.MULTILINE)),
        "num_classes": len(re.findall(r"^class\s+(\w+)", texto, re.MULTILINE)),
    }

def handle_image(path):
    try:
        img = Image.open(str(path))
        return pytesseract.image_to_string(img, config=TESS_CONFIG), {"metodo":"ocr_png", "tamanho": f"{img.width}x{img.height}"}
    except Exception as e:
        return "", {"metodo":"img_err","erro":str(e)}

def handle_json_file(path):
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        texto = json.dumps(data, ensure_ascii=False, indent=2)[:10000]
        return texto, {"metodo":"json","tipo_raiz":type(data).__name__,"dados":data}
    except Exception as e:
        return "", {"metodo":"json_err","erro":str(e)}

HANDLERS_NAO_PDF = {
    ".docx": handle_docx if DOCX_OK else None,
    ".xlsx": handle_xlsx if XLSX_OK else None,
    ".xls": handle_xls if XLS_OK else None,
    ".html": handle_html if BS4_OK else None, ".htm": handle_html if BS4_OK else None,
    ".md": handle_md,
    ".py": handle_py,
    ".png": handle_image if OCR_OK else None, ".jpg": handle_image if OCR_OK else None,
    ".json": handle_json_file,
}
IGNORAR = {".db",".ini",".lnk",".tmp"}
IGNORAR_NAMES = {"Thumbs.db",".DS_Store","desktop.ini"}

# ============================================================
# PIPELINE EXECUTION
# ============================================================
def processar_arquivo(path, categoria, parser_fn, out_folder):
    json_out = out_folder / (safe_filename(path.stem) + (".json" if path.suffix.lower() == ".pdf" else f"{path.suffix.replace('.','_')}.json"))
    if json_out.exists() and json_out.stat().st_size > 500:
        return {"ok": True, "skip": True, "filename": path.name}

    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            texto, meta = extrair_pymupdf(path)
            precisa_ocr = not meta.get("tem_texto")
            campos = parser_fn(texto, path.name) if texto else {}
        elif ext == ".csv":
            with open(path, encoding="utf-8", errors="ignore") as f:
                try: rows = list(csv.DictReader(f, delimiter=";"))
                except: rows = []
            if not rows:
                with open(path, encoding="utf-8", errors="ignore") as f:
                    rows = list(csv.DictReader(f))
            data = {"filename":path.name,"tipo_doc":"CSV_SPCF","categoria":categoria,
                    "total_linhas":len(rows),"campos":list(rows[0].keys()) if rows else [],
                    "dados_completos":rows,"amostra_5_linhas":rows[:5]}
            json_out.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
            return {"ok":True,"filename":path.name}
        elif ext in HANDLERS_NAO_PDF and HANDLERS_NAO_PDF[ext]:
            texto, meta = HANDLERS_NAO_PDF[ext](path)
            precisa_ocr = False
            campos = parse_comum(texto) if texto else {}
        elif ext in IGNORAR or path.name in IGNORAR_NAMES:
            return {"ok":True,"skip":True,"filename":path.name}
        else:
            return {"ok":False,"filename":path.name,"erro":f"ext não suportada: {ext}"}

        data = {
            "filename":path.name,
            "path":str(path.relative_to(ROOT)),
            "categoria":categoria,
            "extensao":ext,
            "data_extracao":date.today().isoformat(),
            "tamanho_bytes":path.stat().st_size,
            "texto_length":len(texto),
            "texto_preview":texto[:800],
            "texto_completo":texto,
            "pdf_meta":meta,
            "precisa_ocr":precisa_ocr if ext == ".pdf" else False,
            "campos_estruturados":campos,
        }
        json_out.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
        return {"ok":True,"filename":path.name,"precisa_ocr":data["precisa_ocr"],"paginas":meta.get("paginas",0)}
    except Exception as e:
        return {"ok":False,"filename":path.name,"erro":str(e)[:200]}

def fase1_fast(src_folder: Path, out: Path, workers=8) -> dict:
    tarefas = []
    for cat_folder in PARSERS.keys():
        folder = src_folder / cat_folder
        if not folder.exists(): continue
        out_cat = out / cat_folder
        out_cat.mkdir(exist_ok=True)
        for f in folder.iterdir():
            if f.is_file() and (f.suffix.lower() == ".pdf" or f.suffix.lower() == ".csv" or f.suffix.lower() in HANDLERS_NAO_PDF):
                tarefas.append((f, cat_folder, PARSERS[cat_folder], out_cat))
    csv_folder = src_folder / "07_SCRAPING_SPCF"
    if csv_folder.exists():
        out_csv = out / "07_SCRAPING_SPCF"
        out_csv.mkdir(exist_ok=True)
        for f in csv_folder.iterdir():
            if f.is_file() and f.suffix.lower() in (".csv",".json"):
                tarefas.append((f, "07_SCRAPING_SPCF", None, out_csv))

    novas = [t for t in tarefas if not (t[3] / (safe_filename(t[0].stem)+".json")).exists()]
    print(f"  Total: {len(tarefas)} | Já processadas: {len(tarefas)-len(novas)} | Pendentes: {len(novas)}")

    ini = time.time()
    results = []
    with ThreadPoolExecutor(max_workers=workers) as ex:
        futures = {ex.submit(processar_arquivo, *t): t for t in novas}
        for i, fut in enumerate(as_completed(futures), 1):
            r = fut.result()
            results.append(r)
            if i % 25 == 0 or i == len(novas):
                print(f"    [{i}/{len(novas)}] {time.time()-ini:.0f}s")
    return {"total":len(tarefas), "novos":len(novas), "sucesso":sum(1 for r in results if r.get("ok")),
            "precisam_ocr": sum(1 for r in results if r.get("precisa_ocr")), "tempo":time.time()-ini}

def fase2_ocr(out: Path, src_folder: Path, workers=6) -> dict:
    pendentes = []
    for folder in out.iterdir():
        if not folder.is_dir(): continue
        for jf in folder.glob("*.json"):
            try: d = json.loads(jf.read_text(encoding="utf-8"))
            except: continue
            if d.get("precisa_ocr") and d.get("texto_length",0) < 100:
                pdf = src_folder / d["categoria"] / d["filename"]
                if pdf.exists():
                    pendentes.append({"json":jf,"pdf":pdf,"data":d})

    def ocr_work(item):
        try:
            texto, meta = ocr_pdf(item["pdf"])
            item["data"]["texto_completo"] = texto
            item["data"]["texto_preview"] = texto[:800]
            item["data"]["texto_length"] = len(texto)
            item["data"]["precisa_ocr"] = False
            item["data"]["pdf_meta"]["metodo"] = "ocr_tesseract"
            item["data"]["pdf_meta"]["paginas_ocr"] = meta.get("paginas_ocr",0)
            item["data"]["campos_estruturados"].update(parse_comum(texto))
            item["json"].write_text(json.dumps(item["data"], ensure_ascii=False, indent=2, default=str), encoding="utf-8")
            return {"ok":True}
        except Exception as e:
            return {"ok":False,"erro":str(e)[:150]}

    if not pendentes:
        return {"pendentes":0,"sucesso":0}
    print(f"  OCR pendentes: {len(pendentes)}")
    ini = time.time()
    ok = 0
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for i, r in enumerate(as_completed(ex.submit(ocr_work, p) for p in pendentes), 1):
            if r.result().get("ok"): ok += 1
            if i % 10 == 0: print(f"    OCR [{i}/{len(pendentes)}] {time.time()-ini:.0f}s | OK={ok}")
    return {"pendentes":len(pendentes),"sucesso":ok,"tempo":time.time()-ini}

def fase3_ingest_contratos(out: Path, orgao: str) -> dict:
    """Ingere contratos extraídos em spcf_contratos."""
    json_dir = out / "01_CONTRATOS"
    if not json_dir.exists():
        return {"inseridos":0,"razao":"sem JSONs de contratos"}

    agregado = defaultdict(lambda: {"arquivos":[],"paginas":0,"valores":set(),"cnpjs":set(),"datas":set(),"reajustes":[],"nes":set(),"clausulas_11":[],"texto_principal":""})
    for jf in json_dir.glob("*.json"):
        try: d = json.loads(jf.read_text(encoding="utf-8"))
        except: continue
        fn = d.get("filename","")
        m = RE_CONTRATO_FN.search(fn)
        if not m: continue
        num, ano = m.groups()
        if len(ano) == 2: ano = "20" + ano
        ref = f"{int(num):03d}/{ano}"
        a = agregado[ref]
        a["arquivos"].append(fn)
        a["paginas"] += d.get("pdf_meta",{}).get("paginas",0)
        cp = d.get("campos_estruturados",{}) or {}
        for v in (cp.get("valores_brl") or [])[:10]: a["valores"].add(v)
        for c in (cp.get("cnpjs") or []): a["cnpjs"].add(c)
        for dt in (cp.get("datas_encontradas") or []): a["datas"].add(dt)
        a["reajustes"].extend(cp.get("reajustes_pct") or [])
        for ne in (cp.get("nes_referenciadas") or []): a["nes"].add(ne)
        if cp.get("clausula_11"): a["clausulas_11"].append(cp["clausula_11"])
        if cp.get("tipo_doc") == "CONTRATO_PRINCIPAL" and not a["texto_principal"]:
            a["texto_principal"] = (d.get("texto_preview") or "")[:600]

    if not DB_PATH.exists():
        return {"inseridos":0,"razao":"prodam.db ausente"}

    conn = sqlite3.connect(str(DB_PATH))
    cur = conn.cursor()
    # Existentes
    rows = cur.execute("""SELECT json_extract(dados_base,'$.N° do contrato'), id FROM spcf_contratos
                          WHERE UPPER(json_extract(dados_base,'$.Cliente')) LIKE ? OR UPPER(cliente) LIKE ?""",
                       (f"%{orgao.upper()}%", f"%{orgao.upper()}%")).fetchall()
    existentes = set()
    for num, _ in rows:
        m = re.search(r"(\d{1,3})/(\d{2,4})", num or "")
        if m:
            n, a = m.groups()
            if len(a) == 2: a = "20" + a
            existentes.add(f"{int(n):03d}/{a}")

    inseridos = 0
    for ref, info in sorted(agregado.items()):
        if ref == "DESCONHECIDO" or ref in existentes: continue
        pk = f"PDF_{orgao.upper().replace('/','_')}_{ref.replace('/','_')}"
        # Analise índices do texto das cláusulas 11
        indice = "INDETERMINADO"
        texto_todo = "\n".join(info["clausulas_11"])
        idxs = [k for k,rx in RE_INDICES.items() if rx.search(texto_todo)]
        if idxs:
            # Prioriza IGPM se aparecer junto com outros (regime legado)
            if "IGPM" in idxs or "IGP-M" in idxs: indice = "IGPM"
            elif "IPCA" in idxs: indice = "IPCA"
            else: indice = idxs[0]
        # valor máximo plausível
        valores_num = []
        for v in info["valores"]:
            try:
                vn = float(str(v).replace(".","").replace(",","."))
                if 10_000 <= vn <= 100_000_000: valores_num.append(vn)
            except: pass
        valor_max = max(valores_num) if valores_num else None
        datas_ord = sorted(info["datas"])
        dados_base = {
            "Cliente": orgao.upper(),
            "N° do contrato": ref,
            "Valor(R$)": f"{valor_max:,.2f}" if valor_max else "",
            "Fim da vigência": datas_ord[-1] if datas_ord else None,
            "Descrição": info["texto_principal"][:300] or f"Contrato {orgao} {ref}",
            "Status": "ENCERRADO_ARQUIVADO",
            "Origem": f"PDF_INGESTAO_{date.today().isoformat()}",
            "Arquivos_fonte": len(info["arquivos"]),
            "Paginas_total": info["paginas"],
            "Indice_correcao": indice,
            "Reajustes_aplicados": info["reajustes"],
            "CNPJs_detectados": sorted(info["cnpjs"]),
            "NEs_referenciadas": list(info["nes"])[:30],
        }
        detalhes = {"arquivos":info["arquivos"],"num_aditivos":sum(1 for f in info["arquivos"] if re.search(r"TA|aditivo",f,re.IGNORECASE))}
        try:
            cur.execute("INSERT OR REPLACE INTO spcf_contratos (id,numero,cliente,dados_base,detalhes,tem_tramite_pdf) VALUES (?,?,?,?,?,?)",
                        (pk, ref, orgao.upper(), json.dumps(dados_base,ensure_ascii=False,default=str), json.dumps(detalhes,ensure_ascii=False,default=str), 0))
            inseridos += 1
        except Exception as e:
            print(f"    ERR {ref}: {e}")
    conn.commit(); conn.close()
    return {"inseridos":inseridos,"ja_existentes":len(existentes),"total_agregado":len(agregado)}

# ============================================================
# MAIN
# ============================================================
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--orgao", help="Sigla do órgão (SEDUC, DETRAN, SES/SUSAM...)")
    ap.add_argument("--listar", action="store_true", help="Lista órgãos disponíveis")
    ap.add_argument("--pular-ocr", action="store_true")
    ap.add_argument("--pular-ingest", action="store_true")
    ap.add_argument("--so-fase", help="ex: '1,3' — só rodar fases específicas")
    args = ap.parse_args()

    if args.listar:
        for o in listar_orgaos(): print(f"  {o}")
        return

    if not args.orgao:
        print("Uso: --orgao SIGLA | --listar"); return

    src = localizar_pasta_consolidado(args.orgao)
    if not src:
        print(f"❌ Pasta consolidada não encontrada para '{args.orgao}'"); return

    print(f"\n{'='*70}")
    print(f"  PIPELINE COMPLETA — {args.orgao}")
    print(f"  Fonte: {src}")
    print(f"{'='*70}\n")

    out = ROOT / f"{sigla_to_folder(args.orgao)}_CONSOLIDADO_JSON"
    out.mkdir(exist_ok=True)

    fases = {1,2,3} if not args.so_fase else {int(x) for x in args.so_fase.split(",")}

    resumo = {"orgao": args.orgao, "pasta_src": str(src), "pasta_out": str(out)}

    if 1 in fases:
        print("[FASE 1] Fast-path PyMuPDF + handlers não-PDF")
        r = fase1_fast(src, out, workers=8)
        resumo["fase1"] = r
        print(f"  Sucesso: {r['sucesso']}/{r['novos']} | Precisam OCR: {r['precisam_ocr']} | Tempo: {r['tempo']:.0f}s\n")

    if 2 in fases and not args.pular_ocr:
        print("[FASE 2] OCR Tesseract para PDFs imagem")
        r = fase2_ocr(out, src, workers=6)
        resumo["fase2"] = r
        if r["pendentes"]:
            print(f"  OCR: {r['sucesso']}/{r['pendentes']} | Tempo: {r.get('tempo',0):.0f}s\n")
        else:
            print("  Nenhum PDF precisa OCR.\n")

    if 3 in fases and not args.pular_ingest:
        print("[FASE 3] Ingestão de contratos em spcf_contratos")
        r = fase3_ingest_contratos(out, args.orgao)
        resumo["fase3"] = r
        print(f"  Inseridos: {r['inseridos']} contratos\n")

    # MASTER
    all_jsons = []
    for folder in out.iterdir():
        if folder.is_dir():
            for jf in folder.glob("*.json"):
                try: all_jsons.append(json.loads(jf.read_text(encoding="utf-8")))
                except: pass
    por_cat = Counter(d.get("categoria","?") for d in all_jsons)
    por_ext = Counter(d.get("extensao") or Path(d.get("filename","")).suffix.lower() for d in all_jsons)
    por_tipo = Counter(d.get("campos_estruturados",{}).get("tipo_doc") or d.get("tipo_doc","?") for d in all_jsons)
    idx_agg = Counter()
    for d in all_jsons:
        for idx in d.get("campos_estruturados",{}).get("indices_mencionados",[]): idx_agg[idx] += 1

    master = {
        "orgao": args.orgao,
        "data": date.today().isoformat(),
        "total_arquivos": len(all_jsons),
        "por_categoria": dict(por_cat),
        "por_extensao": dict(por_ext),
        "por_tipo_doc": dict(por_tipo),
        "indices_agregados": dict(idx_agg),
        "resumo_pipeline": resumo,
    }
    json.dump(master, open(out / f"MASTER_{sigla_to_folder(args.orgao)}.json", "w", encoding="utf-8"), ensure_ascii=False, indent=2, default=str)

    print(f"\n{'='*70}")
    print(f"  RESUMO FINAL — {args.orgao}")
    print(f"{'='*70}")
    print(f"  Total JSONs: {len(all_jsons)}")
    print(f"  Por categoria: {dict(por_cat)}")
    print(f"  Índices detectados: {dict(idx_agg)}")
    print(f"  Master: {out / f'MASTER_{sigla_to_folder(args.orgao)}.json'}")

if __name__ == "__main__":
    main()
