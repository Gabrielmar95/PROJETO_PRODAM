"""
Gera relatório técnico da pasta-mãe PROJETO_PRODAM em .docx.
Design no padrão PRODAM (azul #1F3864 + dourado #B8963E).
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT = Path("/sessions/gallant-focused-brahmagupta/mnt/PROJETO_PRODAM/DIAGNOSTICO_PASTA_MAE.docx")
PROFILE = Path("/sessions/gallant-focused-brahmagupta/mnt/outputs/profile_pasta_mae.json")
INSIGHTS = Path("/sessions/gallant-focused-brahmagupta/mnt/outputs/insights_consolidados.json")

AZUL = RGBColor(0x1F, 0x38, 0x64)
DOURADO = RGBColor(0xB8, 0x96, 0x3E)
CARVAO = RGBColor(0x2D, 0x2D, 0x2D)
CINZA = RGBColor(0x77, 0x77, 0x77)

data = json.loads(PROFILE.read_text(encoding="utf-8"))
insights = json.loads(INSIGHTS.read_text(encoding="utf-8"))


def brl(v):
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = AZUL
    r.font.name = "Calibri"
    # linha dourada
    p2 = doc.add_paragraph("─" * 50)
    p2.runs[0].font.color.rgb = DOURADO
    p2.runs[0].font.size = Pt(10)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = AZUL
    r.font.name = "Calibri"


def add_p(doc, text, bold=False, italic=False, size=11, color=CARVAO):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid"
    for i, (k, v) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.font.bold = True
        r1.font.color.rgb = AZUL
        r1.font.size = Pt(10)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(str(v))
        r2.font.size = Pt(10)
        r2.font.color.rgb = CARVAO


def add_data_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid"
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        _shade(c, "1F3864")
    # Body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = CARVAO
            # zebra
            if ri % 2 == 0:
                _shade(c, "F2F2F2")
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[col_idx].width = Cm(w)


# ============ BUILD ============
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---- CAPA ----
add_p(doc, "")
add_p(doc, "")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DIAGNÓSTICO TÉCNICO")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = DOURADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PASTA-MÃE PROJETO_PRODAM")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Análise de Cientista de Dados")
r.font.size = Pt(14)
r.font.italic = True
r.font.color.rgb = CINZA

add_p(doc, "")
add_p(doc, "")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─" * 40)
r.font.color.rgb = DOURADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Gerado em: {data['gerado_em']}")
r.font.size = Pt(11)
r.font.color.rgb = CARVAO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Contrato 002/2026 — PRODAM S.A. × Brandão Ozores Advogados")
r.font.size = Pt(11)
r.font.color.rgb = CARVAO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gabriel Mar Sociedade Individual de Advocacia — OAB/AM 15.697")
r.font.size = Pt(10)
r.font.italic = True
r.font.color.rgb = CINZA

doc.add_page_break()

# ---- 1. SUMÁRIO EXECUTIVO ----
add_h1(doc, "1. Sumário Executivo")

pf = data["profiles_json"]

add_p(doc,
    "A pasta PROJETO_PRODAM funciona como data lake jurídico estruturado, "
    "desenhado para alimentar 78 pastas-filhas (uma por devedor), cada uma com "
    "nove subpastas padronizadas (de 01_CONTRATOS a 09_RELATORIOS). Este relatório "
    "analisa a pasta mãe sob a ótica de um cientista de dados: arquitetura, "
    "qualidade, cobertura, padrões replicáveis e gaps críticos."
)

add_h2(doc, "Números-chave")

# Calcular métricas
n_dev_pastas = data["n_devedores_pastas"]
n_dev_profiles = pf.get("n_devedores", 0)
divergencia = n_dev_pastas - n_dev_profiles
valor_exig = 93632770.88
valor_atual = 160936948.41

total_registros_db = sum(t.get("count", 0) or 0 for t in data["prodam_db"]["tabelas"].values())

add_kv_table(doc, [
    ("Pasta-mãe (raiz)", str(data["pasta_raiz"])),
    ("Devedores com pasta _CONSOLIDADO", f"{n_dev_pastas}"),
    ("Devedores em profiles.json (SSOT)", f"{n_dev_profiles}"),
    ("Divergência estrutural", f"{divergencia} pastas sem correspondência no SSOT"),
    ("Valor exigível (total portfólio)", brl(valor_exig)),
    ("Valor atualizado (total portfólio)", brl(valor_atual)),
    ("Base prodam.db (tamanho)", f"{data['prodam_db']['tamanho_mb']} MB"),
    ("Base prodam.db (registros)", f"{total_registros_db:,}".replace(",", ".")),
    ("Base OCR_CONSOLIDADO.db (tamanho)", "579 MB — 24 tabelas"),
    ("Benchmark DETRAN (arquivos)", "455 no _CONSOLIDADO | cobertura 88,9%"),
])

add_h2(doc, "Três descobertas críticas")

add_p(doc,
    "1. Sucesso comprovado mas isolado. Apenas o DETRAN atinge cobertura documental ≥ 80% "
    "(88,9%). Os outros 77 devedores estão em estágios iniciais — 59% têm cobertura entre 20% e 49%, "
    "e 33% têm menos de 20% de cobertura. A pasta mãe nasceu como projeto-piloto DETRAN e só "
    "começou a escalar.", size=11
)

add_p(doc,
    "2. Divergência SSOT × Pastas. CLAUDE.md reporta 71 devedores; profiles.json tem 70; "
    "há 78 pastas _CONSOLIDADO no disco. Existem 8 pastas sem correspondência no SSOT — "
    "possivelmente pastas-fantasma de iterações anteriores. Isso quebra a premissa de "
    f"'profiles.json é a SSOT' fixada na regra #15 do CLAUDE.md.", size=11
)

add_p(doc,
    "3. Pipeline pronta, execução pendente. Os scripts orgao_pipeline_completa.py (605 linhas) "
    "e auditoria_completude_devedor.py (453 linhas) implementam uma pipeline parametrizada "
    "('--orgao SIGLA') aplicável a qualquer devedor. O gargalo não é técnico — "
    "é execução: precisa rodar a pipeline para cada um dos 77 devedores restantes.", size=11
)

doc.add_page_break()

# ---- 2. ARQUITETURA DA PASTA MÃE ----
add_h1(doc, "2. Arquitetura da Pasta-Mãe")

add_p(doc,
    "A pasta PROJETO_PRODAM segue um padrão de data lake com cinco camadas "
    "lógicas. Entender essas camadas é pré-requisito para qualquer análise ou "
    "automação."
)

add_h2(doc, "2.1 Camadas da pasta-mãe")

add_data_table(doc,
    ["Camada", "Local", "Finalidade"],
    [
        ("BRUTO", "SPCF_EXTRACAO/", "Documentos originais raspados do SPCF (HTML, PDF, OCR)"),
        ("ESTRUTURADO", "prodam.db, OCR_CONSOLIDADO.db", "Dados normalizados em SQLite"),
        ("SSOT", "PRODAM_DOCS/profiles.json", "Fonte autoritativa única de verdade sobre devedores"),
        ("PASTA-FILHA", "PRODAM_DOCS/<DEV>_CONSOLIDADO/", "Uma por devedor; 9 subpastas padronizadas"),
        ("ENTREGÁVEL", "PRODAM_DOCS/<DEV>_DOSSIE/", "Dossiês jurídicos finais (dashboard, CSV, LLM)"),
    ],
    col_widths=[2.5, 5.5, 8]
)

add_h2(doc, "2.2 Schema padronizado das pastas-filhas")

add_p(doc,
    "Toda pasta-filha _CONSOLIDADO segue o mesmo schema de 9 subpastas. Isso é o "
    "que torna a pasta mãe 'replicável': ao adicionar um novo devedor, basta "
    "seguir o molde."
)

add_data_table(doc,
    ["Subpasta", "Conteúdo", "% devedores com arquivos"],
    [
        ("01_CONTRATOS", "Contratos originais + Termos Aditivos (PDFs)", f"{100 - 92.3:.1f}% (apenas {78 - 72} devedores)"),
        ("02_EMPENHOS", "Notas de Empenho (NE) — Art. 202 VI CC", f"{100 - 96.2:.1f}% (apenas {78 - 75})"),
        ("03_FATURAS", "Notas Fiscais emitidas pela PRODAM", f"{100 - 74.4:.1f}% ({78 - 58})"),
        ("04_NOTAS_LIQUIDACAO", "Liquidação formal (Lei 4.320/64)", f"{100 - 96.2:.1f}% ({78 - 75})"),
        ("05_ACEITES", "Aceites técnicos do fiscal", f"{100 - 38.5:.1f}% ({78 - 30})"),
        ("06_COBRANCAS", "Cartas e ofícios de cobrança", f"{100 - 19.2:.1f}% ({78 - 15})"),
        ("07_SCRAPING_SPCF", "Dados raspados do sistema SPCF", f"{100 - 98.7:.1f}% (apenas 1)"),
        ("08_PDFS_ORIGINAIS", "PDFs brutos antes de catalogação", f"{100 - 47.4:.1f}% ({78 - 37})"),
        ("09_RELATORIOS", "Relatórios consolidados do devedor", f"{100 - 98.7:.1f}% (apenas 1)"),
    ],
    col_widths=[3.5, 6.5, 4]
)

add_p(doc, "",)
add_p(doc,
    "Leitura crítica: três subpastas têm cobertura < 10% (07_SCRAPING_SPCF, "
    "09_RELATORIOS, 02_EMPENHOS, 04_NOTAS_LIQUIDACAO). Isso é grave porque "
    "empenhos (NE) e notas de liquidação (NL) são exatamente o material que "
    "interrompe prescrição (Art. 202 VI CC) e confirma serviço prestado (Lei 4.320/64). "
    "Sem esses documentos, a força probatória cai para FRACA.",
    italic=True, size=10, color=CINZA
)

doc.add_page_break()

# ---- 3. QUALIDADE DE DADOS ----
add_h1(doc, "3. Qualidade de Dados")

add_h2(doc, "3.1 Base prodam.db (91,5 MB, SQLite)")

tabelas_db = []
for nome, info in data["prodam_db"]["tabelas"].items():
    if nome == "sqlite_sequence":
        continue
    tabelas_db.append((nome, f"{info['count']:,}".replace(",", "."), str(info["cols"])))

add_data_table(doc, ["Tabela", "Registros", "Colunas"], tabelas_db,
               col_widths=[5, 3, 2])

add_p(doc, "")
add_p(doc,
    f"Total de registros: {total_registros_db:,}".replace(",", ".") +
    f". Views: {', '.join(data['prodam_db']['views'])}.",
    size=10
)

add_h2(doc, "3.2 profiles.json (SSOT de 70 devedores)")

add_kv_table(doc, [
    ("Total devedores", "70"),
    ("Campos por devedor (sample SES/SUSAM)", "~120 campos"),
    ("Tem CNPJ preenchido", "69/70 (98,6%)"),
    ("Tem data_prescricao_proxima", "52/70 (74,3%)"),
    ("Tem score_composto", "51/70 (72,9%)"),
    ("Tem p_recuperacao", "51/70 (72,9%)"),
    ("Título executivo = True", "2/70 (2,9%)"),
    ("Faturas totais (soma)", "3.497"),
    ("Faturas prescritas", "1.085 (31,0%)"),
    ("Faturas exigíveis", "2.412 (69,0%)"),
])

add_h2(doc, "3.3 Distribuição por categoria")

add_data_table(doc,
    ["Categoria", "Devedores", "% do portfólio"],
    [
        ("GOV_DIRETA (secretarias)", "26", "37,1%"),
        ("GOV_INDIRETA (autarquias, fundações)", "21", "30,0%"),
        ("EMPRESA_PRIVADA", "22", "31,4%"),
        ("(sem classificação)", "1", "1,4%"),
    ],
    col_widths=[6, 3, 3]
)

add_h2(doc, "3.4 Distribuição por força probatória")

add_data_table(doc,
    ["Força", "Devedores", "Tratamento jurídico"],
    [
        ("FORTE", "12 (17,1%)", "Execução direta (Art. 784 CPC) — cadeia completa"),
        ("MÉDIA", "15 (21,4%)", "Ação monitória (Art. 700 CPC) — cadeia parcial"),
        ("FRACA", "42 (60,0%)", "Cobrança simples ou negociação — docs incompletos"),
        ("(vazio)", "1 (1,4%)", "Pendente classificação"),
    ],
    col_widths=[2.5, 3, 8]
)

add_h2(doc, "3.5 Divergências entre fontes")

add_data_table(doc,
    ["Fonte", "Devedores", "Observação"],
    [
        ("CLAUDE.md", "71", "Métrica declarada no topo do arquivo"),
        ("profiles.json (SSOT)", "70", "Fonte oficial"),
        ("Pastas _CONSOLIDADO", "78", "+ 8 divergentes"),
        ("Pastas _DOSSIE", "82", "+ 12 (inclui DESCONHECIDO, MASTER, SEJEL, SRMM sem CONSOLIDADO)"),
        ("Tabela devedores (prodam.db)", "81", "+ 11 em relação ao SSOT"),
    ],
    col_widths=[5, 2.5, 6]
)

add_p(doc, "")
add_p(doc,
    "Recomendação: executar auditoria de reconciliação entre as 5 fontes para "
    "descobrir se as pastas extras são devedores legítimos (e devem entrar no SSOT) "
    "ou resíduos de iterações passadas (e devem ir para _legado/).",
    italic=True, size=10, color=CINZA
)

doc.add_page_break()

# ---- 4. BENCHMARK DETRAN ----
add_h1(doc, "4. Benchmark: DETRAN como Pasta-Filha Ideal")

add_p(doc,
    "A pasta DETRAN_CONSOLIDADO é a única que atinge padrão A+ (score 94/100 "
    "conforme CLAUDE.md). Serve como referência para replicação. Análise detalhada:"
)

detran_cons = data["cobertura_consolidado"]["DETRAN"]
rows_detran = []
for sub, info in detran_cons["subpastas"].items():
    if info.get("missing"):
        rows_detran.append((sub, "0", "0,00 MB", "FALTANDO"))
    else:
        rows_detran.append((
            sub,
            f"{info.get('files', 0):,}".replace(",", "."),
            f"{info.get('size_mb', 0):.2f} MB",
            "OK" if info.get("files", 0) > 0 else "VAZIO",
        ))

add_data_table(doc, ["Subpasta", "Arquivos", "Tamanho", "Status"], rows_detran,
               col_widths=[4.5, 2.5, 2.5, 2.5])

add_p(doc, "")
add_p(doc,
    "Mesmo o DETRAN (benchmark) tem 04_NOTAS_LIQUIDACAO vazia. A ausência de NLs "
    "é padrão do portfólio — provavelmente porque a Administração Pública não "
    "costuma fornecer cópia da NL ao contratado. Alternativa: usar NE (empenho) "
    "+ NF atestada como substituto funcional.",
    italic=True, size=10, color=CINZA
)

add_h2(doc, "Gap por volume de arquivos")

add_p(doc, "Se o DETRAN atinge A+ com 455 arquivos, os demais devedores precisam "
    "atingir volumes comparáveis. Situação atual dos top-10 por valor exigível:")

# Top 10 por valor (via CLAUDE.md — SEDUC, SES, SSP, SEJUSC, SEAD, BRADESCO, CETAM, SEDECTI, SALUX, PC)
cobertura_cons = data["cobertura_consolidado"]
top_devedores_clauded = ["SEDUC", "SES", "SSP", "SEJUSC", "SEAD", "CETAM", "SEDECTI", "POLICIA_CIVIL"]

rows_top = [("DETRAN (benchmark)", "455", "88,9%", "GAP DE REFERÊNCIA: 0")]
for dev in top_devedores_clauded:
    info = cobertura_cons.get(dev, {})
    files = info.get("total_arquivos", 0)
    cov = info.get("cobertura_pct", 0)
    gap = 455 - files
    rows_top.append((dev, f"{files:,}".replace(",", "."), f"{cov}%", f"{gap:+,}".replace(",", ".")))

add_data_table(doc, ["Devedor", "Arquivos", "Cobertura", "Gap vs DETRAN"],
               rows_top, col_widths=[4.5, 2.5, 2.5, 3])

doc.add_page_break()

# ---- 5. PADRÕES REPLICÁVEIS ----
add_h1(doc, "5. Padrões Replicáveis (Scripts e Automação)")

add_p(doc,
    "A pasta mãe contém uma infraestrutura técnica madura. Mapeei os scripts "
    "que são 'moldes' para replicar o processamento a qualquer devedor."
)

add_h2(doc, "5.1 Scripts-molde identificados")

add_data_table(doc,
    ["Script", "Linhas", "Função"],
    [
        ("orgao_pipeline_completa.py", "605", "Pipeline genérica parametrizada por --orgao SIGLA"),
        ("auditoria_completude_devedor.py", "453", "Checklist de 11 itens por devedor"),
        ("dossie_multiformato_devedor.py", "420", "Gera 5 formatos por devedor"),
        ("reconciliacao_4_fontes.py", "667", "Reconcilia profiles × prodam.db × SPCF × pastas"),
        ("prodam_utils.py", "212", "Utilidades: norm(), brl(), datas, parse"),
        ("auto_update_claude_md.py", "283", "Regenera CLAUDE.md automaticamente"),
        ("analisar_v4.py", "1.775", "Análise forense (o maior — audita tudo)"),
    ],
    col_widths=[5.5, 2, 7]
)

add_h2(doc, "5.2 Duplicações detectadas (risco de drift)")

add_p(doc,
    "Scripts que existem em duas localizações podem divergir silenciosamente. "
    "Verificação por hash MD5:"
)

add_data_table(doc,
    ["Arquivo", "Status", "Ação recomendada"],
    [
        ("auto_update_claude_md.py", "IDÊNTICO", "Manter. Considerar link simbólico"),
        ("prodam_utils.py", "IDÊNTICO", "Manter. Considerar link simbólico"),
        ("consultas.py", "DIVERGENTE", "Unificar versão ativa; legado → _legado/"),
        ("orgao_pipeline_completa.py", "DIVERGENTE", "Consolidar — 605 vs 606 linhas"),
    ],
    col_widths=[5, 3, 7]
)

add_h2(doc, "5.3 Infraestrutura jurídica (REFERENCIA_JURIDICA)")

add_p(doc,
    "Árvore de 18 subpastas (00_KNOWLEDGE_BASE a 17_ERRATA_CORRECOES) "
    "que funciona como enciclopédia jurídica do projeto. É a matéria-prima "
    "para fundamentar cada peça processual. A estrutura está madura; falta "
    "cruzamento com os devedores (qual norma aplica-se a qual caso)."
)

doc.add_page_break()

# ---- 6. RANKING COMPLETO ----
add_h1(doc, "6. Ranking Completo — 78 Devedores")

add_p(doc,
    "Cada devedor classificado por volume de arquivos (proxy de maturidade). "
    "Quanto mais abaixo, mais trabalho de coleta documental ainda falta."
)

ranking = insights["ranking_arquivos"]
rows_rank = []
for i, item in enumerate(ranking, 1):
    rows_rank.append((
        str(i),
        item["devedor"],
        f"{item['arquivos']:,}".replace(",", "."),
        f"{item['cobertura']}%",
    ))

# Quebrar em páginas — max 35 por página
for i in range(0, len(rows_rank), 35):
    chunk = rows_rank[i:i + 35]
    add_data_table(doc, ["#", "Devedor", "Arquivos", "Cobertura"], chunk,
                   col_widths=[1.5, 5.5, 3, 3])
    if i + 35 < len(rows_rank):
        doc.add_page_break()

doc.add_page_break()

# ---- 7. RECOMENDAÇÕES ----
add_h1(doc, "7. Recomendações — Roadmap do Cientista de Dados")

add_h2(doc, "7.1 Ações imediatas (esta semana)")

add_p(doc, "1. Reconciliar SSOT × pastas. Rodar reconciliacao_4_fontes.py e decidir "
    "o destino das 8 pastas _CONSOLIDADO sem correspondência no profiles.json. "
    "Opções: entrar no SSOT ou ir para _legado/.")

add_p(doc, "2. Unificar duplicações de scripts. Eleger versão canônica entre "
    "PRODAM_DOCS/ e scripts/ para consultas.py e orgao_pipeline_completa.py. "
    "Demais (auto_update, prodam_utils) estão idênticos mas é frágil — trocar por link simbólico.")

# Corrigido em 2026-05-10 via cascata profiles.json — fundamento PASSO6 (Of. 129/2021 reclassificado: ato do credor não interrompe prescrição)
add_p(doc, "3. Monitorar SES/SUSAM. Próxima prescrição em 2026-08-31 "
    "(status MEDIA, D+113 base 10/05/2026; reclassificação PASSO6 — Of. 129/2021 não interrompe). "
    "py -3.12 orgao_pipeline_completa.py --orgao \"SES/SUSAM\".")

add_h2(doc, "7.2 Ações de médio prazo (próximo mês)")

add_p(doc, "4. Escalar pipeline a TOP-10 por valor. SEDUC (R$ 49M), SES (R$ 14M), "
    "SSP (R$ 4,5M), SEJUSC, SEAD, CETAM, SEDECTI, POLICIA_CIVIL — rodar "
    "orgao_pipeline_completa.py para cada um. Cada execução sobe a cobertura do ~40% para ~90%.")

add_p(doc, "5. Atacar subpastas vazias. 02_EMPENHOS e 04_NOTAS_LIQUIDACAO estão "
    "vazias em 96% dos devedores. Usar baixar_lacunas_spcf.py para automatizar "
    "o download via SPCF (requer VPN PRODAM).")

add_p(doc, "6. Validar 23 prescritos. 23 devedores estão classificados como "
    "'PRESCRITA' em urgencia_prescricao. Antes de baixar esses créditos, rodar "
    "analise-prescricao-creditos para confirmar que não há marco interruptivo "
    "(Art. 202 VI CC) que zere a contagem.")

add_h2(doc, "7.3 Ações estratégicas (trimestre)")

add_p(doc, "7. Profissionalizar o pipeline como produto. A pasta mãe é um "
    "framework jurídico genérico (não específico do PRODAM). Com pequenos ajustes, "
    "serve para qualquer cobrança B2B contra administração pública. Considerar "
    "empacotá-la como produto/serviço replicável a outros clientes do escritório.")

add_p(doc, "8. Adicionar observabilidade. A pipeline roda mas não deixa rastro "
    "padronizado. Adicionar logging estruturado (JSON Lines) em cada execução "
    "para poder auditar o quê/quando/quanto foi processado.")

add_p(doc, "9. Reduzir duplicação de dados. OCR_CONSOLIDADO.db (579 MB) + prodam.db (91 MB) "
    "+ pastas com PDFs brutos + SPCF_EXTRACAO — há sobreposição. Mapear se dá "
    "para consolidar ou se cada um tem razão própria de existir.")

doc.add_page_break()

# ---- APÊNDICE ----
add_h1(doc, "Apêndice A — Glossário técnico")

add_data_table(doc,
    ["Termo", "Significado"],
    [
        ("SSOT", "Single Source of Truth — fonte única da verdade. Aqui, profiles.json"),
        ("Pasta-mãe", "PROJETO_PRODAM — raiz de tudo, contém infraestrutura + devedores"),
        ("Pasta-filha", "Pasta de um devedor específico (<DEV>_CONSOLIDADO + <DEV>_DOSSIE)"),
        ("Data lake", "Repositório que guarda dados em múltiplos formatos (raw + estruturado)"),
        ("Cobertura", "% das 9 subpastas do _CONSOLIDADO que contêm arquivos"),
        ("SPCF", "Sistema interno da PRODAM — fonte primária de empenhos e faturas"),
        ("OCR", "Optical Character Recognition — leitura automática de PDFs de imagem"),
        ("Benchmark", "Referência de excelência. Aqui, DETRAN com score 94/100"),
    ],
    col_widths=[3.5, 10]
)

add_h2(doc, "Apêndice B — Arquivos desta análise")

add_data_table(doc,
    ["Arquivo", "Conteúdo"],
    [
        ("DIAGNOSTICO_PASTA_MAE.docx", "Este documento (relatório técnico formatado)"),
        ("DIAGNOSTICO_PASTA_MAE.html", "Dashboard interativo equivalente"),
        ("profile_pasta_mae.json", "JSON cru com todos os dados coletados"),
        ("insights_consolidados.json", "Agregações e rankings"),
        ("insights_devedores.csv", "Tabela pivô: devedor × cobertura × arquivos"),
        ("profiles_resumo.csv", "Resumo do profiles.json por devedor"),
        ("profile_pasta_mae.py", "Script que coleta os dados (fonte reproduzível)"),
    ],
    col_widths=[5.5, 8]
)

add_p(doc, "")
add_p(doc, "")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Fim do relatório —")
r.font.italic = True
r.font.color.rgb = CINZA
r.font.size = Pt(10)

# Salvar
doc.save(str(OUT))
print(f"[OK] Relatório salvo em: {OUT}")
print(f"Tamanho: {OUT.stat().st_size / 1024:.1f} KB")
