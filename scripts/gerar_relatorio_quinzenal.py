"""
gerar_relatorio_quinzenal.py — Gera o Relatório Quinzenal à PRODAM (MD + DOCX).

Contrato 002/2026 — PRODAM S.A. × Brandão Ozores Advogados.
Elaborado por Gabriel Mar (OAB/AM 15.697).

Uso (Windows):
  py -3.12 scripts\\gerar_relatorio_quinzenal.py ^
      --periodo-inicio 2026-05-28 --periodo-fim 2026-06-11 ^
      --numero "[005/2026]" ^
      --out DOCUMENTOS_GERADOS/RELATORIOS_QUINZENAIS ^
      [--profiles CAMINHO/profiles.json]

Regras observadas:
  - Valores BRL SEMPRE em Decimal (prodam_utils.brl / fmt_brl) — nunca float.
  - Datas ISO no código; formato brasileiro (DD/MM/AAAA) no texto do relatório.
  - Anti-alucinação: a Seção 4 (Ações Realizadas) contém apenas fatos com lastro
    em arquivos do repositório, citados entre parênteses. O que não tem registro
    é marcado "[a confirmar com o advogado]".
  - Numeração e período ficam entre colchetes para conferência pelo advogado.
  - Sem nota "Elaboração técnica" (decisão institucional permanente de 29/04/2026,
    relatorios/CHANGELOG_SESSAO_2026-04-29_design_brandao_ozores.md).

KPIs: o script soma o profiles.json informado (via --profiles). Se os totais
divergirem materialmente da referência oficial (CLAUDE.md / STATUS_DEVEDORES.md,
regenerados em 10/06/2026 17:45 a partir do SSOT restaurado), o relatório usa a
referência oficial no corpo e registra a divergência em nota de rodapé — caso
típico quando o único profiles.json disponível no clone é um snapshot antigo.

Design DOCX (Brandão Ozores): azul #1F3864, dourado #B8963E, Calibri 11,
tabelas com cabeçalho azul e texto branco. Helpers adaptados de
scripts/gerar_relatorio_docx.py (que é one-shot, com paths fixos — por isso
este gerador é autocontido).
"""
from __future__ import annotations

import argparse
import sys
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from prodam_utils import brl, fmt_brl, load_profiles  # noqa: E402

# ============================================================
# REFERÊNCIA OFICIAL — CLAUDE.md + STATUS_DEVEDORES.md de 10/06/2026 17:45
# (satélites regenerados via auto_update_claude_md.py a partir do SSOT
#  PRODAM_DOCS/profiles.json restaurado em 10/06/2026).
# Atualizar este bloco a cada nova regeneração dos satélites.
# ============================================================
REF = {
    "fonte": "CLAUDE.md e STATUS_DEVEDORES.md regenerados em 10/06/2026 17:45 a partir do SSOT (PRODAM_DOCS/profiles.json)",
    "n_devedores": 69,
    "val_exig": brl("83668078.44"),
    "val_atualizado": brl("125245390.64"),
    "n_com_atualizado": 65,
    "faturas_total": 3477,
    "faturas_exigiveis": 2326,
    "faturas_prescritas": 1082,
    "faturas_fora_universo": 69,
    "categorias": [("Governo — Adm. Direta", 26), ("Governo — Adm. Indireta", 21), ("Empresas Privadas", 22)],
    "forca": [("FORTE", 12), ("MÉDIA", 15), ("FRACA", 42)],
    "fases": [("F0 (diagnóstico/documentação)", 51), ("F0_DIAGNOSTICO", 4), ("F3 (TRD)", 9), ("F5 (petição pronta)", 5)],
    "pipeline": [
        ("ANALISAR_DOCUMENTACAO", 36), ("CLASSIFICAR", 17), ("ENVIAR_TRD", 9),
        ("PROTOCOLAR_PETICAO", 5), ("AVALIAR_SUCESSAO", 1), ("HABILITAÇÃO DE CRÉDITO", 1),
    ],
    # (sigla, exigível, atualizado, força, próximo passo, observação prescrição)
    "top5": [
        ("SEDUC", brl("49215512.48"), brl("50263263.56"), "FORTE", "ANALISAR_DOCUMENTACAO", "751 dias"),
        ("SES/SUSAM", brl("4783356.52"), brl("8230061.40"), "FORTE", "ENVIAR_TRD", "112 dias"),
        ("SSP", brl("4553230.80"), brl("29034062.63"), "FORTE", "PROTOCOLAR_PETICAO", "🔴 30/06/2026"),
        ("SEJUSC", brl("2589660.12"), brl("5262537.82"), "MÉDIA", "ANALISAR_DOCUMENTACAO", "🟡 31/08/2026"),
        ("SEAD", brl("2339702.20"), brl("4296278.76"), "FORTE", "ENVIAR_TRD", "477 dias"),
    ],
    "alertas": [
        ("🔴 SSP", "30/06/2026", brl("4553230.80"),
         "Petição pronta (fase F5). Protocolar antes do prazo, com blindagem pré-execução, checklist de 20 itens e verificação das exceções do Decreto Estadual AM nº 53.464/2026 (art. 1º, §§1º–4º)."),
        ("🔴 SUHAB", "30/06/2026", brl("840061.15"),
         "Em ANALISAR_DOCUMENTACAO. Acelerar análise e adotar medida interruptiva da prescrição antes do prazo."),
        ("🟡 SEJUSC", "31/08/2026", brl("2589660.12"),
         "Concluir análise documental até o prazo."),
        ("🛡️ DETRAN", "marco interruptivo", Decimal(0),
         "Protegido por marco interruptivo (Art. 202, VI, CC) com data registrada no passado — marco prevalece; reconfirmar e atualizar a data. Monitorar cutoff da NF 110654 (CT 179/2018): 19/08/2026."),
    ],
    "nota_alertas": "9 devedores sem data de prescrição registrada (AADESAM, BRADESCO, CASA MILITAR, CGE, FJJA, FMPES, SEJEL, SETRAB, UGPI) — levantamento de vencimentos pendente.",
}

# ============================================================
# SEÇÃO 4 — AÇÕES REALIZADAS NA QUINZENA (editar a cada relatório).
# Somente fatos com lastro em arquivos do repositório (fonte entre parênteses).
# ============================================================
ACOES_QUINZENA = {
    "4.1 Gestão do portfólio e prescrição": [
        "09/06 — Diagnóstico integral do portfólio (somente leitura): mapeamento pronto × pendente por devedor, "
        "identificação de 14 devedores acionáveis (5 com petição pronta, 9 com TRD pronto/recomendado), gaps de "
        "reconciliação SPCF × exigível e priorização P1–P7 com a SSP em primeiro lugar "
        "(fonte: registros internos de trabalho, sessão de 09/06).",
        "09/06 — Recálculo das datas de prescrição no profiles.json: cerca de 30 devedores saíram da data-placeholder "
        "2026-03-20 para datas individualizadas, com backup pré-recálculo preservado; a metodologia do recálculo segue "
        "em conferência (fonte: registros internos de trabalho, sessão de 10/06).",
        "10/06 — Panorama do portfólio com números conferidos entre profiles.json, TASKS.md e STATUS_DEVEDORES.md; "
        "geração de dashboard HTML autocontido e de minuta de relatório quinzenal em .docx (documento de gestão, "
        "datado de 10/06); o envio desse documento à PRODAM está [a confirmar com o advogado] "
        "(fonte: registros internos de trabalho, sessão de 10/06).",
        "10/06 — Programação de 3 alertas pontuais de prazo: 24/06 (prescrição SSP + SUHAB em 30/06), 01/08 (SEJUSC "
        "em 31/08) e 05/08 (DETRAN — marco interruptivo e cutoff NF 110654 em 19/08) "
        "(fonte: registros internos de trabalho, sessão de 10/06).",
    ],
    "4.2 Integridade de dados e acervo documental": [
        "28/05 — Saneamento de qualidade de código e dados (3 commits): eliminação de usos de float em valores "
        "monetários (regra Decimal), consolidação dos helpers de formatação BRL e auditoria forense da reversão de "
        "drift de 13/05 no profiles.json, com 113 testes automatizados aprovados "
        "(fonte: registros internos de trabalho, sessão de 28/05).",
        "10/06 — Detecção e restauração do profiles.json (SSOT) corrompido/truncado: restaurado a partir de snapshot "
        "íntegro de 09/06, com backup do arquivo corrompido preservado e validação pós-restauração (69 devedores; "
        "soma R$ 83.668.078,44); satélites do projeto (CLAUDE.md, STATUS_DEVEDORES.md) regenerados às 17:45 "
        "(fonte: registros internos de trabalho, sessão de 10/06).",
        "10/06 — Organização não-destrutiva do acervo completo do projeto: 91.256 arquivos inventariados (32,4 GB), "
        "70.298 cópias organizadas por devedor/tipo com verificação de integridade MD5 em 100% das cópias, 20.958 "
        "duplicatas exatas identificadas (não recopiadas) e manifesto auditável de 91.256 linhas; originais intactos "
        "(fonte: registros internos de trabalho, relatório de organização de 10/06).",
        "10/06 — DETRAN (Sessão 2): renomeação forense de 229 CSVs do subprojeto DETRAN_AUDITORIA_COMPLETA com hash "
        "MD5 pré-ação, índice de auditoria (CSV+JSON) e rollback integral possível; nenhum PDF tocado; alterações "
        "sincronizadas com o GitHub (commit 031b20a0) "
        "(fonte: registros internos de trabalho, sessão de 10/06).",
        "09/06 — Organização da pasta \"Arquivos Principais Projeto PRODAM\" em modo cópia (16 arquivos copiados com "
        "conferência MD5; 3 duplicatas detectadas), com identificação documental relevante: o PDF \"Canal Comunicação "
        "Oficial\" é, na verdade, o Ofício PRESI/045/2026 da PRODAM (resposta ao Ofício 001/2026), reclassificado "
        "para a pasta de ofícios (fonte: registros internos de trabalho, sessão de 09/06).",
    ],
    "4.3 Governança e controles internos": [
        "08–10/06 — Aprimoramento contínuo das ferramentas internas de auditoria, organização documental e controle "
        "de prazos do escritório aplicadas ao portfólio (auditoria do ferramental, automação de painéis de "
        "acompanhamento e rotinas de verificação de integridade de dados) "
        "(fonte: registros internos de trabalho, sessões de 08 a 10/06).",
    ],
    "4.4 Em andamento nesta data (11/06/2026)": [
        "SEDUC — Memorial preliminar de cálculo em elaboração sobre o universo recente do SPCF (106 faturas; base "
        "R$ 54.535.717,29), com atualização monetária pela SELIC/EC 113 até a competência 04/2026 e ressalvas "
        "expressas de regime presumido (índice contratual por contrato ainda em confirmação). O universo SPCF "
        "recente supera o snapshot consolidado de mar/2026 (84 faturas; R$ 49.215.512,48 exigíveis) por critérios "
        "de corte distintos — conciliação detalhada na Seção 2 do memorial.",
        "SEDUC — Plano de auditoria documental do acervo (DPCON, pendrive e SPCF) em elaboração, para suprir a "
        "cadeia probatória (Contrato + NE + NF + Atesto) do maior crédito do portfólio.",
    ],
    "4.5 Registro negativo (transparência)": [
        "Não há, nos arquivos do repositório relativos ao período, registro de protocolos judiciais, notificações "
        "extrajudiciais enviadas a devedores ou reuniões externas nesta quinzena [a confirmar com o advogado].",
    ],
}

PROXIMOS_PASSOS = [
    ("SSP — até 30/06/2026 (crítico)",
     "Protocolar a petição pronta antes da prescrição de 30/06/2026 (R$ 4.553.230,80 em risco): rodar blindagem "
     "pré-execução, checklist de 20 itens e verificar as exceções do Decreto Estadual AM nº 53.464/2026 "
     "(art. 1º, §§1º–4º)."),
    ("SUHAB — até 30/06/2026 (crítico)",
     "Acelerar a análise documental e adotar medida interruptiva da prescrição antes de 30/06/2026 "
     "(R$ 840.061,15 em risco)."),
    ("SEDUC — maior crédito (R$ 49,2 mi)",
     "Finalizar o memorial preliminar de cálculo (106 faturas SPCF; SELIC/EC 113 até 04/2026) e executar o plano de "
     "auditoria documental do acervo (DPCON/pendrive/SPCF), inclusive reconciliação da divergência SPCF × exigível "
     "apontada no diagnóstico de 09/06."),
    ("DETRAN — destravar protocolo",
     "Concluir a revisão interna da notificação extrajudicial (crédito apurado de R$ 28.196.572,22) e requisitar à "
     "PRODAM, via LAI/ofício, os contratos-base faltantes; monitorar o marco de 19/08/2026."),
    ("SES/SUSAM e SEAD — TRDs",
     "Emitir e enviar os Termos de Reconhecimento de Dívida (SES/SUSAM: R$ 4.783.356,52, prescrição em ~112 dias; "
     "SEAD: R$ 2.339.702,20)."),
    ("SEJUSC — até 31/08/2026",
     "Concluir a análise documental (R$ 2.589.660,12)."),
    ("Higiene de dados",
     "Levantar vencimentos dos 9 devedores sem data de prescrição registrada; conferir a metodologia do recálculo de "
     "prescrição de 09/06; investigar a causa do truncamento do profiles.json e adotar escrita atômica com validação."),
    ("Obrigação contratual",
     "Manter a cadência dos relatórios quinzenais à PRODAM (multa contratual de R$ 500,00/dia de atraso)."),
]

REFERENCIAS = [
    "Base consolidada de devedores do projeto (SSOT), restaurada e validada em 10/06/2026.",
    "Painéis consolidados do portfólio regenerados em 10/06/2026 17:45 (KPIs, alertas de prescrição e pipeline).",
    "Registros internos de trabalho do período 28/05–11/06/2026 — disponíveis para auditoria da PRODAM mediante solicitação.",
    "Memorial preliminar de cálculo SEDUC (em elaboração, 11/06/2026).",
]


# ============================================================
# Datas
# ============================================================
def mil(n) -> str:
    """3477 -> '3.477' (milhar pt-BR)."""
    return f"{int(n):,}".replace(",", ".")


def iso_para_br(s: str) -> str:
    """'2026-05-28' → '28/05/2026'."""
    return datetime.strptime(s, "%Y-%m-%d").strftime("%d/%m/%Y")


# ============================================================
# KPIs do profiles.json (conferência contra a referência oficial)
# ============================================================
def calcular_kpis_profiles(path: str | Path) -> dict:
    p = load_profiles(path)
    tot_exig = Decimal(0)
    tot_atu = Decimal(0)
    ft = fe = fp = 0
    for d in p.values():
        tot_exig += brl(d.get("val_exig"))
        tot_atu += brl(d.get("val_atualizado"))
        ft += int(d.get("faturas_total") or 0)
        fe += int(d.get("faturas_exigiveis") or 0)
        fp += int(d.get("faturas_prescritas") or 0)
    return {
        "n_devedores": len(p),
        "val_exig": tot_exig,
        "val_atualizado": tot_atu,
        "faturas_total": ft,
        "faturas_exigiveis": fe,
        "faturas_prescritas": fp,
    }


def diverge(kpis: dict) -> bool:
    if kpis["n_devedores"] != REF["n_devedores"]:
        return True
    for campo in ("val_exig", "val_atualizado"):
        a, b = kpis[campo], REF[campo]
        maior = max(abs(a), abs(b))
        if maior and abs(a - b) / maior > Decimal("0.005"):
            return True
    return False


# ============================================================
# Modelo do documento (blocos) — renderizado em MD e DOCX
# ============================================================
def montar_blocos(numero: str, ini_iso: str, fim_iso: str, data_emissao_iso: str,
                  nota_rodape: str | None) -> list:
    ini_br, fim_br = iso_para_br(ini_iso), iso_para_br(fim_iso)
    emissao_br = iso_para_br(data_emissao_iso)
    B: list = []

    B.append(("titulo", f"Relatório Quinzenal nº {numero} — Contrato 002/2026"))
    B.append(("subtitulo", "Recuperação de Créditos — PRODAM S.A. × Brandão Ozores Advogados"))
    B.append(("meta", [
        ("Cliente", "PRODAM — Processamento de Dados Amazonas S.A. (CNPJ 04.407.920/0001-80)"),
        ("Contratante", "Brandão Ozores Advogados"),
        ("Elaboração", "Gabriel Mar — OAB/AM 15.697"),
        ("Período de referência", f"[{ini_br} a {fim_br}]"),
        ("Data de emissão", emissao_br),
    ]))

    # ---- 1. Sumário Executivo ----
    B.append(("h1", "1. Sumário Executivo"))
    B.append(("p",
              f"O portfólio sob gestão compreende {REF['n_devedores']} devedores, com crédito exigível de "
              f"{fmt_brl(REF['val_exig'])} e valor atualizado de {fmt_brl(REF['val_atualizado'])} "
              f"({REF['n_com_atualizado']}/{REF['n_devedores']} devedores com atualização monetária calculada). "
              f"O universo de faturas soma {mil(REF['faturas_total'])} documentos: {mil(REF['faturas_exigiveis'])} exigíveis, "
              f"{mil(REF['faturas_prescritas'])} prescritas e {REF['faturas_fora_universo']} fora do universo de cobrança "
              f"(canceladas/excluídas). (*)"))
    B.append(("p",
              "Destaques da quinzena: (i) restauração do arquivo-mestre de dados do portfólio (profiles.json), "
              "detectado corrompido em 10/06, com validação integral pós-restauração; (ii) recálculo das datas de "
              "prescrição, individualizando prazos antes registrados com data-placeholder; (iii) organização "
              "não-destrutiva de todo o acervo documental (70.298 cópias verificadas por hash MD5); e "
              "(iv) início do memorial preliminar de cálculo da SEDUC, maior crédito do portfólio "
              "(R$ 49,2 milhões exigíveis), em andamento nesta data."))
    B.append(("p",
              "Pontos de atenção imediata: SSP e SUHAB têm prescrição projetada para 30/06/2026 "
              "(seção 3). Não houve recuperação de valores no período."))

    # ---- 2. Panorama do Portfólio ----
    B.append(("h1", "2. Panorama do Portfólio"))
    B.append(("h2", "2.1 Composição por categoria, força probatória e fase"))
    B.append(("tabela", ["Categoria", "Devedores"], [(c, str(n)) for c, n in REF["categorias"]]))
    B.append(("tabela", ["Força probatória", "Devedores"], [(c, str(n)) for c, n in REF["forca"]]))
    B.append(("tabela", ["Fase do pipeline", "Devedores"], [(c, str(n)) for c, n in REF["fases"]]))
    B.append(("h2", "2.2 Pipeline — próximo passo por devedor"))
    B.append(("tabela", ["Próximo passo", "Devedores"], [(c, str(n)) for c, n in REF["pipeline"]]))
    B.append(("h2", "2.3 Cinco maiores devedores (por valor exigível)"))
    B.append(("tabela",
              ["Devedor", "Exigível", "Atualizado", "Força", "Próximo passo", "Prescrição"],
              [(sig, fmt_brl(e), fmt_brl(a), f, pp, presc) for sig, e, a, f, pp, presc in REF["top5"]]))
    B.append(("p", "Lista completa dos 69 devedores em STATUS_DEVEDORES.md (anexo de referência, seção 6)."))

    # ---- 3. Alertas de Prescrição ----
    B.append(("h1", "3. Alertas de Prescrição"))
    B.append(("p",
              "A prescrição é apurada por fatura individual, contada do vencimento (arts. 189 e 206, §5º, I, do "
              "Código Civil). Os marcos abaixo são as datas críticas registradas na base do projeto:"))
    B.append(("tabela_alerta",
              ["Devedor", "Data crítica", "Valor exigível", "Providência"],
              [(dev, data, fmt_brl(v) if v else "—", prov) for dev, data, v, prov in REF["alertas"]]))
    B.append(("p_destaque",
              "ATENÇÃO: SSP (R$ 4.553.230,80) e SUHAB (R$ 840.061,15) prescrevem em 30/06/2026 — "
              "ação interruptiva impreterível dentro do mês corrente."))
    B.append(("p", REF["nota_alertas"]))

    # ---- 4. Ações Realizadas na Quinzena ----
    B.append(("h1", "4. Ações Realizadas na Quinzena"))
    B.append(("p",
              f"Período de referência: [{ini_br} a {fim_br}]. Relacionam-se apenas fatos com registro nos arquivos "
              "do projeto (fonte indicada entre parênteses); itens sem registro documental estão marcados "
              "\"[a confirmar com o advogado]\"."))
    for sub, itens in ACOES_QUINZENA.items():
        B.append(("h2", sub))
        for item in itens:
            B.append(("bullet", item))

    # ---- 5. Próximos Passos ----
    B.append(("h1", "5. Próximos Passos"))
    for titulo, texto in PROXIMOS_PASSOS:
        B.append(("bullet_bold", titulo, texto))

    # ---- 6. Anexos e Referências ----
    B.append(("h1", "6. Anexos e Referências"))
    B.append(("p", "Documentos e bases que lastreiam este relatório (disponíveis no repositório do projeto):"))
    for r in REFERENCIAS:
        B.append(("bullet", r))

    if nota_rodape:
        B.append(("h2", "Nota de rodapé"))
        B.append(("nota", "(*) " + nota_rodape))

    B.append(("assinatura", ["Gabriel Mar", "OAB/AM 15.697", "Gabriel Mar Sociedade Individual de Advocacia"]))
    return B


def nota_divergencia(kpis: dict, profiles_path: str) -> str:
    return (
        "Conferência de KPIs: o único profiles.json disponível neste clone do repositório "
        f"({profiles_path}) é um snapshot antigo e diverge da SSOT viva — soma {kpis['n_devedores']} registros, "
        f"{fmt_brl(kpis['val_exig'])} exigível, {fmt_brl(kpis['val_atualizado'])} atualizado e "
        f"{mil(kpis['faturas_total'])} faturas ({mil(kpis['faturas_exigiveis'])} exigíveis / {mil(kpis['faturas_prescritas'])} "
        "prescritas). Os números do corpo deste relatório seguem os painéis oficiais regenerados em 10/06/2026 "
        "17:45 a partir da base consolidada restaurada. O exigível consolidado não inclui o crédito do DETRAN "
        "(R$ 28.196.572,22, apurado em notificação extrajudicial em revisão final), registrado por R$ 0,00 no "
        "consolidado-mãe; com sua inclusão, o portfólio exigível é da ordem de R$ 111,9 milhões. "
        "Reconciliação em curso."
    )


# ============================================================
# Render MD
# ============================================================
def render_md(blocos: list) -> str:
    L: list[str] = []
    for b in blocos:
        t = b[0]
        if t == "titulo":
            L += [f"# {b[1]}", ""]
        elif t == "subtitulo":
            L += [f"## {b[1]}", ""]
        elif t == "meta":
            for k, v in b[1]:
                L.append(f"**{k}:** {v}  ")
            L.append("")
            L += ["---", ""]
        elif t == "h1":
            L += [f"## {b[1]}", ""]
        elif t == "h2":
            L += [f"### {b[1]}", ""]
        elif t in ("p", "nota"):
            L += [b[1], ""]
        elif t == "p_destaque":
            L += [f"> 🔴 **{b[1]}**", ""]
        elif t in ("tabela", "tabela_alerta"):
            headers, rows = b[1], b[2]
            L.append("| " + " | ".join(headers) + " |")
            L.append("|" + "|".join("---" for _ in headers) + "|")
            for row in rows:
                L.append("| " + " | ".join(str(c) for c in row) + " |")
            L.append("")
        elif t == "bullet":
            L.append(f"- {b[1]}")
        elif t == "bullet_bold":
            L.append(f"- **{b[1]}** — {b[2]}")
        elif t == "assinatura":
            L += ["", "---", ""]
            for linha in b[1]:
                L.append(f"**{linha}**  " if linha == b[1][0] else f"{linha}  ")
            L.append("")
    # bullets precisam de linha em branco após blocos consecutivos
    out: list[str] = []
    for i, linha in enumerate(L):
        out.append(linha)
        if linha.startswith("- ") and i + 1 < len(L) and not L[i + 1].startswith("- ") and L[i + 1] != "":
            out.append("")
    return "\n".join(out).rstrip() + "\n"


# ============================================================
# Render DOCX — design Brandão Ozores (#1F3864 / #B8963E / Calibri 11)
# ============================================================
def render_docx(blocos: list, out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    AZUL = RGBColor(0x1F, 0x38, 0x64)
    DOURADO = RGBColor(0xB8, 0x96, 0x3E)
    CARVAO = RGBColor(0x2D, 0x2D, 0x2D)
    CINZA = RGBColor(0x77, 0x77, 0x77)
    VERMELHO = RGBColor(0xB0, 0x00, 0x00)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    def _shade(cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_hex)
        tcPr.append(shd)

    def _run(p, text, size=11, bold=False, italic=False, color=CARVAO):
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        return r

    def linha_dourada():
        p = doc.add_paragraph()
        _run(p, "─" * 50, size=10, color=DOURADO)

    def add_tabela(headers, rows):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            _run(c.paragraphs[0], h, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shade(c, "1F3864")
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = t.rows[ri + 1].cells[ci]
                c.text = ""
                _run(c.paragraphs[0], str(val), size=9)
                if ri % 2 == 0:
                    _shade(c, "F2F2F2")
        doc.add_paragraph()

    for b in blocos:
        t = b[0]
        if t == "titulo":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, b[1].upper(), size=18, bold=True, color=AZUL)
            p.paragraph_format.space_after = Pt(6)
        elif t == "subtitulo":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, b[1], size=12, bold=True, color=DOURADO)
            p.paragraph_format.space_after = Pt(12)
        elif t == "meta":
            for k, v in b[1]:
                p = doc.add_paragraph()
                _run(p, f"{k}: ", size=11, bold=True, color=AZUL)
                _run(p, v, size=11)
            linha_dourada()
        elif t == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(4)
            _run(p, b[1].upper(), size=15, bold=True, color=AZUL)
            linha_dourada()
        elif t == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            _run(p, b[1], size=12, bold=True, color=AZUL)
        elif t == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            _run(p, b[1])
        elif t == "p_destaque":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            _run(p, b[1], size=11, bold=True, color=VERMELHO)
        elif t == "nota":
            p = doc.add_paragraph()
            _run(p, b[1], size=9, italic=True, color=CINZA)
        elif t in ("tabela", "tabela_alerta"):
            add_tabela(b[1], b[2])
        elif t == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _run(p, b[1])
        elif t == "bullet_bold":
            p = doc.add_paragraph(style="List Bullet")
            _run(p, b[1] + " — ", bold=True, color=AZUL)
            _run(p, b[2])
        elif t == "assinatura":
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, "─" * 30, size=10, color=DOURADO)
            for i, linha in enumerate(b[1]):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0:
                    _run(p, linha, size=12, bold=True, color=AZUL)
                else:
                    _run(p, linha, size=10, color=CINZA)
    doc.save(str(out_path))


# ============================================================
# CLI
# ============================================================
def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Gera o Relatório Quinzenal à PRODAM (MD + DOCX).")
    ap.add_argument("--periodo-inicio", required=True, help="Data ISO (ex.: 2026-05-28)")
    ap.add_argument("--periodo-fim", required=True, help="Data ISO (ex.: 2026-06-11)")
    ap.add_argument("--numero", required=True, help='Número entre colchetes (ex.: "[005/2026]")')
    ap.add_argument("--out", required=True, help="Pasta de saída")
    ap.add_argument("--profiles", default=None,
                    help="Caminho do profiles.json para conferência de KPIs (opcional)")
    ap.add_argument("--data-emissao", default=None, help="Data ISO de emissão (default: hoje)")
    args = ap.parse_args(argv)

    data_emissao = args.data_emissao or date.today().isoformat()
    numero = args.numero.strip()
    if not numero.startswith("["):
        numero = f"[{numero.strip('[]')}]"

    nota = None
    if args.profiles:
        kpis = calcular_kpis_profiles(args.profiles)
        if diverge(kpis):
            nota = nota_divergencia(kpis, args.profiles)
            print("[AVISO] profiles.json diverge da referência oficial — nota de rodapé adicionada.")
            print(f"        JSON: {kpis['n_devedores']} devedores · exigível {fmt_brl(kpis['val_exig'])}")
            print(f"        REF : {REF['n_devedores']} devedores · exigível {fmt_brl(REF['val_exig'])}")
        else:
            print("[OK] profiles.json confere com a referência oficial.")

    blocos = montar_blocos(numero, args.periodo_inicio, args.periodo_fim, data_emissao, nota)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    base = f"Relatorio_Quinzenal_PRODAM_{args.periodo_fim}"
    md_path = out_dir / f"{base}.md"
    docx_path = out_dir / f"{base}.docx"

    md_path.write_text(render_md(blocos), encoding="utf-8")
    print(f"[OK] MD  : {md_path}")
    render_docx(blocos, docx_path)
    print(f"[OK] DOCX: {docx_path} ({docx_path.stat().st_size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
