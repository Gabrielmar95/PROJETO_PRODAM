#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera notificação extrajudicial para SES/SUSAM com python-docx"""

import json
from pathlib import Path
from datetime import datetime
from decimal import Decimal
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
if os.environ.get("PRODAM_FREEZE_EMISSAO"):
    import sys; sys.exit("[FREEZE] Emissão de peças bloqueada durante auditoria DE. Remover PRODAM_FREEZE_EMISSAO para destravar.")
# ============================================================
# BLOQUEIO MANUAL — adicionado em 2026-05-10
# Motivo: numeração "NOT/001/2026" hardcoded colide com a NE 001/2026
# DETRAN já protocolada via ICP-Brasil em 20/04/2026.
# Para destravar: substituir todas as ocorrências de "NOT/001/2026"
# por numeração SES não conflitante e remover este bloqueio.
# ============================================================
import sys
sys.exit("[BLOQUEADO] Numeração NOT/001/2026 colide com NE DETRAN. Ver bloco de comentário no topo do arquivo.")

# Caminhos
PROFILES_PATH = Path(r"C:\Users\gabri\Desktop\PROJETO_PRODAM\PRODAM_DOCS\profiles.json")
OUTPUT_DIR = Path(r"C:\Users\gabri\Desktop\PROJETO_PRODAM\PRODAM_DOCS\_SKILLS\dossie-juridico-prodam-workspace\iteration-1\eval-2-notificacao-ses\with_skill\outputs")
OUTPUT_FILE = OUTPUT_DIR / "NOT_SES_SUSAM_001_2026.docx"

# Cores
COR_AZUL = RGBColor(0x1F, 0x38, 0x64)
COR_PRINCIPAL = RGBColor(0x2D, 0x2D, 0x2D)

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Carregar dados
with open(PROFILES_PATH, 'r', encoding='utf-8') as f:
    profiles = json.load(f)
dados = profiles["SES/SUSAM"]

# Preparar valores
val_exig = Decimal(str(dados.get("val_exig", "14748048.96")))
val_exig_fmt = f"R$ {val_exig:,.2f}".replace(",", ".")

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.6)

# Título
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NOTIFICAÇÃO EXTRAJUDICIAL")
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = COR_AZUL
p.paragraph_format.space_after = Pt(12)

# Separador
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("————————————————————")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0xB8, 0x96, 0x3E)
p.paragraph_format.space_after = Pt(12)

# Número
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NOT/001/2026")
r.font.bold = True
p.paragraph_format.space_after = Pt(8)

# Data
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hoje = datetime.now()
data_fmt = hoje.strftime("%d de %B de %Y").replace(
    "January", "janeiro").replace("February", "fevereiro").replace("March", "março").replace(
    "April", "abril").replace("May", "maio").replace("June", "junho").replace(
    "July", "julho").replace("August", "agosto").replace("September", "setembro").replace(
    "October", "outubro").replace("November", "novembro").replace("December", "dezembro")
r = p.add_run(f"Manaus, {data_fmt}.")
p.paragraph_format.space_after = Pt(18)

# Destinatário
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("DESTINATÁRIO")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
r = p.add_run("Secretaria de Saúde do Amazonas / Fundação de Medicina Tropical")
r.font.bold = True
r = p.add_run("\nCNPJ: 00.697.295/0001-05\nManaus, Amazonas")
p.paragraph_format.space_after = Pt(18)

# Assunto
p = doc.add_paragraph()
r = p.add_run("ASSUNTO")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run("Notificação para pagamento de dívida vencida conforme art. 784 do Código de Processo Civil.")
p.paragraph_format.space_after = Pt(18)

# Conteúdo
p = doc.add_paragraph()
r = p.add_run("CONTEÚDO")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

# Parágrafo introdutório
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run("A ")
r = p.add_run("Gabriel Mar Sociedade Individual de Advocacia")
r.font.color.rgb = COR_PRINCIPAL
r = p.add_run(", agindo como representante legal de ")
r = p.add_run("PRODAM S.A.")
r.font.bold = True
r = p.add_run(" (CNPJ ")
r = p.add_run("04.407.920/0001-80")
r.font.bold = True
r = p.add_run("), vem, por este meio, ")

# Notificar
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run("NOTIFICAR")
r.font.bold = True
r = p.add_run(" a ")
r = p.add_run("SES/SUSAM")
r.font.bold = True
r = p.add_run(" do débito em aberto de ")

# Valor
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run(f"{val_exig_fmt} (quatorze milhões, setecentos e quarenta e oito mil, quarenta e oito reais e noventa e seis centavos)")
r.font.bold = True
r = p.add_run(", referente a faturas vencidas não pagas, resultante de contratações legítimas de serviços.")
p.paragraph_format.space_after = Pt(12)

# Fundamentos
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("FUNDAMENTOS JURÍDICOS")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

items = [
    ("I.", "A notificação fundamenta-se no art. 784 do Código de Processo Civil, que admite execução extrajudicial contra devedor que reconhece a dívida, bem como no reconhecimento tácito evidenciado por empenhos, notas de liquidação e aceites técnicos."),
    ("II.", "O valor referido encontra-se devidamente atualizado conforme índice SELIC, em conformidade com a Lei 14.905/2024 e os arts. 404 a 406 do Código Civil."),
    ("III.", "A composição documental (contrato + empenho + nota de liquidação + aceite) constitui título executivo extrajudicial, conforme jurisprudência consolidada (REsp 793.969/RJ, Rel. p/ acórdão Min. José Delgado; Teori Zavascki vencido)."),
]

for letra, texto in items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(letra + " ")
    r.font.bold = True
    r = p.add_run(texto)
    p.paragraph_format.space_after = Pt(6)

# Prazo
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("PRAZO PARA PAGAMENTO")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run("Fica a ")
r = p.add_run("SES/SUSAM")
r.font.bold = True
r = p.add_run(" notificada para que proceda ao pagamento da dívida no prazo de ")
r = p.add_run("15 (quinze) dias úteis")
r.font.bold = True
r = p.add_run(" contados do recebimento desta notificação, sob pena de ajuizamento de ação executiva.")
p.paragraph_format.space_after = Pt(12)

# Alerta prescrição
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("ALERTA — PRESCRIÇÃO IMINENTE")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run("ATENÇÃO: ")
r.font.bold = True
r = p.add_run("O direito de PRODAM prescreve em ")
r = p.add_run("13 de maio de 2026")
r.font.bold = True
r = p.add_run(" (em ")
r = p.add_run("29 dias")
r.font.bold = True
r = p.add_run("). Passada esta data, a cobrança será prejudicada pela fluência do prazo prescricional de 5 (cinco) anos contado do vencimento das faturas (art. 206, §5º, I do CC).")
p.paragraph_format.space_after = Pt(12)

# Consequências
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("CONSEQUÊNCIAS DO NÃO PAGAMENTO")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

conseqs = [
    ("1.", "Sem resposta em 15 dias úteis: ajuizamento de ação executiva perante o Tribunal de Justiça do Estado do Amazonas, com todos os ônus processuais (custas, honorários, multa de 10%)."),
    ("2.", "Perda da prescrição: caso o prazo de prescrição se complete (31/08/2026), PRODAM perderia todo e qualquer direito de cobrança, sem direito a regressão."),
    ("3.", "Regime de execução: como entidade da administração direta, a execução se dará via precatório ou Requisição de Pequeno Valor (RPV), conforme art. 100 CF."),
]

for num, texto in conseqs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(num + " ")
    r.font.bold = True
    r = p.add_run(texto)
    p.paragraph_format.space_after = Pt(6)

# Acordo
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("POSSIBILIDADE DE ACORDO")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run("PRODAM permanece aberta a negociações. Manifestações ou propostas de pagamento devem ser encaminhadas a este escritório no prazo supracitado.")
p.paragraph_format.space_after = Pt(12)

# Foro
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("FORO COMPETENTE")
r.font.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run("Fica reservado o direito de PRODAM de promover execução judicial perante o Tribunal de Justiça do Estado do Amazonas (Foro de Manaus).")
p.paragraph_format.space_after = Pt(24)

# Encerramento
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Respeitosamente apresentado,")
p.paragraph_format.space_after = Pt(24)

# Assinatura
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gabriel Mar")
r.font.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OAB/AM 15.697")
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Advogado")
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gabriel Mar Sociedade Individual de Advocacia")
r.font.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Em nome de Brandão Ozores Advogados")

# Salvar
try:
    doc.save(OUTPUT_FILE)
    print(f"[✓] Notificação gerada com sucesso!")
    print(f"    Arquivo: {OUTPUT_FILE}")
except Exception as e:
    print(f"[✗] Erro ao salvar: {e}")
    import traceback
    traceback.print_exc()
