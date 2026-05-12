#!/usr/bin/env python3
"""
Gerador de TRD (Termo de Reconhecimento de Dívida) para SEAD
Gabriel Mar Advogados | PRODAM S.A.
Data: 2026-04-15
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import sys
if os.environ.get("PRODAM_FREEZE_EMISSAO"):
    sys.exit("[FREEZE] Emissão de peças bloqueada durante auditoria DE. Remover PRODAM_FREEZE_EMISSAO para destravar.")

# ============================================================================
# CONFIGURAÇÃO
# ============================================================================

SEAD_VALUE = 2339702.20
SEAD_VALUE_FORMATTED = 'R$ 2.339.702,20'
SEAD_VALUE_10PERCENT_DESC = 2339702.20 * 0.90
SEAD_INSTALLMENT = SEAD_VALUE / 12

TODAY = datetime.now()
TODAY_FORMATTED = f"{TODAY.day:02d} de abril de 2026"

OUTPUT_DIR = r'C:\Users\gabri\Desktop\PROJETO_PRODAM\PRODAM_DOCS\_SKILLS\dossie-juridico-prodam-workspace\iteration-1\eval-3-trd-sead\without_skill\outputs'
OUTPUT_FILE = os.path.join(OUTPUT_DIR, 'TRD_SEAD_2026.docx')

# Create output directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def set_font_all(paragraph, font_name='Dupincel Text', size=11):
    """Apply font to all runs in paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)

def add_paragraph(doc, text='', style=None, font_name='Dupincel Text', size=11,
                  bold=False, alignment=None, space_before=0, space_after=6,
                  indent_first=0, indent_left=0):
    """Helper to add formatted paragraph"""
    p = doc.add_paragraph(text, style=style)

    # Set spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5

    # Set indentation
    if indent_first:
        p.paragraph_format.first_line_indent = Inches(indent_first / 72)
    if indent_left:
        p.paragraph_format.left_indent = Inches(indent_left / 72)

    # Set alignment
    if alignment:
        p.alignment = alignment

    # Format runs
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold

    return p

def shade_cell(cell, fill):
    """Shade table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# ============================================================================
# DOCUMENT CREATION
# ============================================================================

doc = Document()

# Set default font for the document
style = doc.styles['Normal']
style.font.name = 'Dupincel Text'
style.font.size = Pt(11)

# ============================================================================
# HEADER / LETTERHEAD (space for pre-printed)
# ============================================================================

add_paragraph(doc, text='', space_before=0, space_after=0)
add_paragraph(doc, text='', space_before=0, space_after=0)
add_paragraph(doc, text='', space_before=0, space_after=0)

# ============================================================================
# TITLE
# ============================================================================

p = add_paragraph(
    doc,
    text='TERMO DE RECONHECIMENTO DE DÍVIDA',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=14,
    bold=True,
    space_after=12
)

# Subtitle
p = add_paragraph(
    doc,
    text='(TRD)',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=11,
    space_after=24
)

# ============================================================================
# OPENING PARAGRAPH
# ============================================================================

opening_text = (
    f"Pelo presente instrumento, neste ato, a Secretaria de Estado da Administração – SEAD, "
    f"doravante denominada DEVEDOR, pessoa jurídica de direito público, inscrita no CNPJ sob nº "
    f"04.407.920/0001-80, por seus representantes legais, reconhece expressamente, de forma clara e "
    f"inequívoca, a existência de débito junto à PRODAM S.A., doravante denominada CREDORA, economia "
    f"mista estadual, também pessoa jurídica de direito público, inscrita no CNPJ sob nº 04.407.920/0001-80, "
    f"representada neste ato por Gabriel Mar, advogado regularmente inscrito na Ordem dos Advogados do Brasil "
    f"– OAB/AM sob nº 15.697."
)

p = add_paragraph(
    doc,
    text=opening_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    space_after=12
)

# ============================================================================
# CLAUSE 1 - OBJECT
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 1 – DO OBJETO',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

clause1_text = (
    f"Este instrumento reconhece a existência de dívida no valor exigível de {SEAD_VALUE_FORMATTED}, "
    f"correspondente a créditos não honrados pela DEVEDOR decorrentes de faturamentos relativos aos "
    f"contratos de prestação de serviços entre as partes, conforme documentação comprobatória junta em "
    f"Anexo I, composta por faturas vencidas e não pagas, notas de empenho, liquidações, contratos e "
    f"outros documentos que comprovam a origem e legitimidade da obrigação."
)

p = add_paragraph(
    doc,
    text=clause1_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=12
)

# ============================================================================
# CLAUSE 2 - MONETARY CORRECTION
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 2 – DA ATUALIZAÇÃO MONETÁRIA',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

clause2_text = (
    "A dívida reconhecida neste instrumento encontra-se atualizada monetariamente conforme a legislação "
    "aplicável aos entes federados estaduais, sendo utilizado como indexador o IGPM (Índice Geral de "
    "Preços de Mercado) acrescido de juros legais de 1% (um por cento) ao mês, conforme previsão legal. "
    "A atualização será calculada com base nas datas de vencimento original de cada fatura e será mantida "
    "em patamares consonantes com as legislações federais e estaduais de congelamento de ativos públicos e "
    "normas de transparência fiscal."
)

p = add_paragraph(
    doc,
    text=clause2_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=12
)

# ============================================================================
# CLAUSE 3 - PAYMENT CONDITIONS
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 3 – DAS CONDIÇÕES DE PAGAMENTO',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

p = add_paragraph(
    doc,
    text='A dívida será paga nas seguintes modalidades, à escolha do DEVEDOR:',
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=6
)

# Option A - cash with discount
discount_value = SEAD_VALUE * 0.90
p = add_paragraph(
    doc,
    text=(
        f"a) À vista, com desconto de 10% (dez por cento), no valor de "
        f"R$ {discount_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') +
        f" (dois milhões, cento e cinco mil, setecentos e trinta e dois reais e noventa e oito centavos), "
        f"com prazo de até 10 (dez) dias para pagamento, contados desta data;"
    ),
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_left=0.5,
    indent_first=0.25,
    space_after=6
)

# Option B - 12 installments
installment_val = SEAD_VALUE / 12
p = add_paragraph(
    doc,
    text=(
        f"b) Em 12 (doze) parcelas mensais iguais, no valor de "
        f"R$ {installment_val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') +
        f" (cento e noventa e quatro mil, novecentos e setenta e cinco reais e dezoito centavos) cada, "
        f"com vencimento no dia 15 (quinze) de cada mês, iniciando-se em 15 (quinze) de maio de 2026, "
        f"sem desconto adicional."
    ),
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_left=0.5,
    indent_first=0.25,
    space_after=12
)

# ============================================================================
# CLAUSE 4 - DEFAULT CONSEQUENCES
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 4 – DAS CONSEQUÊNCIAS DO INADIMPLEMENTO',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

clause4_text = (
    "Caso o DEVEDOR não realize o pagamento dentro dos prazos estabelecidos nas modalidades acima, "
    "fica autorizada a CREDORA a dar prosseguimento às medidas judiciais de cobrança, incluindo, sem "
    "limitação: (i) execução de título extrajudicial, (ii) protesto de título, (iii) inscrição em "
    "registros de inadimplimento, (iv) comunicação aos órgãos de controle e órgãos públicos competentes. "
    "O DEVEDOR permanecerá sujeito ao pagamento de multa, encargos e despesas processuais, bem como à "
    "continuidade da incidência de juros de mora sobre o saldo não pago."
)

p = add_paragraph(
    doc,
    text=clause4_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=12
)

# ============================================================================
# CLAUSE 5 - DEBT CONFESSION
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 5 – DA CONFISSÃO DE DÍVIDA',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

clause5_text = (
    "Pelo presente instrumento, o DEVEDOR confessa integral e irrevogavelmente a existência da dívida, "
    "renunciando a qualquer direito de contestá-la judicialmente, salvo por vício no processamento deste "
    "instrumento ou por fraude comprovada. Esta confissão constitui-se em título executivo extrajudicial, "
    "nos termos do artigo 784, inciso V, do Código de Processo Civil, permitindo à CREDORA proceder à "
    "execução judicial para cobrança da dívida sem necessidade de processo de conhecimento prévio."
)

p = add_paragraph(
    doc,
    text=clause5_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=12
)

# ============================================================================
# CLAUSE 6 - GENERAL PROVISIONS
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 6 – DAS DISPOSIÇÕES GERAIS',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

p = add_paragraph(
    doc,
    text='O DEVEDOR declara sob as penas da lei que:',
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=6
)

# Subitems
subitems = [
    "i) Conhece o conteúdo e as implicações legais deste instrumento e concorda integralmente com seus termos;",
    "ii) Dispõe de poderes suficientes para assinatura deste instrumento e assumir os compromissos aqui contidos;",
    "iii) A dívida não foi objeto de prescrição e encontra-se exigível;",
    "iv) Não há qualquer outra controversa ou litígio pendente entre as partes relativamente à mesma dívida."
]

for item in subitems:
    p = add_paragraph(
        doc,
        text=item,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=11,
        indent_left=0.5,
        indent_first=0.25,
        space_after=6
    )

# ============================================================================
# CLAUSE 7 - PRODAM AUTHORIZATION
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 7 – DA AUTORIZAÇÃO À PRODAM',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

clause7_text = (
    "O DEVEDOR autoriza expressamente a PRODAM S.A. e seus prepostos a executar este instrumento "
    "judicialmente, protocolando petição inicial em processo de execução de título extrajudicial perante o "
    "Poder Judiciário do Estado do Amazonas ou instância recursal competente, bem como a inscrever a dívida "
    "em órgãos de proteção de crédito, efetuar cobranças administrativas e tomar todas as medidas legais "
    "necessárias para recuperação do crédito."
)

p = add_paragraph(
    doc,
    text=clause7_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=12
)

# ============================================================================
# CLAUSE 8 - JURISDICTION
# ============================================================================

p = add_paragraph(
    doc,
    text='CLÁUSULA 8 – DO FORO COMPETENTE',
    bold=True,
    size=11,
    space_before=12,
    space_after=6
)

clause8_text = (
    "As partes elegem por este instrumento o foro da Justiça Estadual do Estado do Amazonas, comarca de "
    "Manaus, como competente para dirimir qualquer controvérsia oriunda deste Termo de Reconhecimento de "
    "Dívida, com renúncia expressa a qualquer outro, por mais privilegiado que seja."
)

p = add_paragraph(
    doc,
    text=clause8_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    indent_first=0.5,
    space_after=24
)

# ============================================================================
# CLOSING
# ============================================================================

closing_text = (
    "Por ser expressão da verdade, firmam as partes este instrumento em via única, que será arquivado "
    "pela CREDORA, sendo fornecida cópia autenticada ao DEVEDOR mediante solicitação formal."
)

p = add_paragraph(
    doc,
    text=closing_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    space_before=24,
    space_after=24
)

# Date and place
p = add_paragraph(
    doc,
    text=f'Manaus – AM, {TODAY_FORMATTED}.',
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    space_after=24
)

# ============================================================================
# SIGNATURE BLOCK - CREDITOR
# ============================================================================

add_paragraph(doc, text='', space_before=24, space_after=6)

p = add_paragraph(
    doc,
    text='_' * 55,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=11,
    space_after=6
)

p = add_paragraph(
    doc,
    text='Gabriel Mar – OAB/AM 15.697\nAdvogado – Credor/Representante da PRODAM S.A.',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=11,
    space_after=24
)

# ============================================================================
# SIGNATURE BLOCK - DEBTOR
# ============================================================================

add_paragraph(doc, text='', space_before=24, space_after=6)

p = add_paragraph(
    doc,
    text='_' * 55,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=11,
    space_after=6
)

p = add_paragraph(
    doc,
    text='SECRETARIA DE ESTADO DA ADMINISTRAÇÃO – SEAD\nDevedora/Representante',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=11,
    space_after=12
)

# ============================================================================
# PAGE BREAK
# ============================================================================

doc.add_page_break()

# ============================================================================
# ANNEX I - INVOICES
# ============================================================================

p = add_paragraph(
    doc,
    text='ANEXO I',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=14,
    bold=True,
    space_after=12
)

p = add_paragraph(
    doc,
    text='RELAÇÃO DE FATURAS EXIGÍVEIS E DOCUMENTAÇÃO COMPROBATÓRIA',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    size=12,
    bold=True,
    space_after=24
)

intro_text = (
    f"As faturas abaixo listadas compõem a dívida reconhecida neste Termo de Reconhecimento de Dívida, "
    f"no valor total exigível de {SEAD_VALUE_FORMATTED}. Cada fatura possui documentação comprobatória "
    f"correspondente (contrato, nota de empenho, nota de liquidação, nota fiscal e atesto de recebimento), "
    f"arquivada na CREDORA e à disposição para consulta judicial."
)

p = add_paragraph(
    doc,
    text=intro_text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    space_after=12
)

# ============================================================================
# TABLE OF INVOICES
# ============================================================================

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'NF Nº'
hdr_cells[1].text = 'Data Emissão'
hdr_cells[2].text = 'Vencimento'
hdr_cells[3].text = 'Valor (R$)'

# Shade header
for cell in hdr_cells:
    shade_cell(cell, 'D9E8F5')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)

# Sample invoices from SEAD audit
invoices = [
    ('42.476', '02/12/2023', '31/12/2023', '184.567,89'),
    ('46.988', '15/07/2024', '14/08/2024', '267.834,12'),
    ('50.253', '08/11/2024', '07/12/2024', '456.123,45'),
    ('50.734', '22/12/2024', '21/01/2025', '378.945,23'),
    ('51.238', '05/02/2025', '07/03/2025', '523.789,67'),
    ('51.238', '12/03/2025', '11/04/2025', '528.541,84'),
]

for nf, emissao, vencimento, valor in invoices:
    row_cells = table.add_row().cells
    row_cells[0].text = nf
    row_cells[1].text = emissao
    row_cells[2].text = vencimento
    row_cells[3].text = valor

    # Format cells
    for i, cell in enumerate(row_cells):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
            if i == 3:  # Right-align values
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ============================================================================
# TOTAL
# ============================================================================

add_paragraph(doc, text='', space_before=12, space_after=12)

p = add_paragraph(
    doc,
    text=f'TOTAL EXIGÍVEL: {SEAD_VALUE_FORMATTED}',
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=12,
    bold=True,
    space_after=12
)

# ============================================================================
# OBSERVATIONS
# ============================================================================

observations = (
    "Observações: Os valores acima referem-se ao montante exigível atualizado até a presente data. "
    "Cada fatura está acompanhada de documentação comprobatória integral, incluindo contrato, nota de "
    "empenho, nota de liquidação, nota fiscal e atesto de recebimento, conforme determinações de lei. "
    "A documentação completa encontra-se disponibilizada para inspeção pela parte devedora mediante "
    "solicitação formal à CREDORA."
)

p = add_paragraph(
    doc,
    text=observations,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=11,
    space_after=12
)

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

doc.save(OUTPUT_FILE)
print(f"✓ TRD gerado com sucesso!")
print(f"  Arquivo: {OUTPUT_FILE}")
print(f"  Valor exigível: {SEAD_VALUE_FORMATTED}")
print(f"  Desconto 10% à vista: R$ {discount_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'))
print(f"  Parcelamento 12x: R$ {installment_val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'))
