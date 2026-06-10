from docx import Document
import re

doc = Document('/home/ubuntu/upload/DOSSIE_COBRANCA_DETRAN_v4.docx')

# Inspect all paragraphs with their styles and text
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "None"
        # Check for section headers (numbered items)
        if re.match(r'^\d+(\.\d+)?\s*[—–-]', text) or style.startswith('Heading'):
            print(f"[{i}] Style: '{style}' | Text: '{text[:120]}'")
        elif i < 5:
            print(f"[{i}] Style: '{style}' | Text: '{text[:120]}'")

print("\n\n--- ALL PARAGRAPHS (first 200 chars) ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "None"
        print(f"[{i}] Style: '{style}' | Text: '{text[:200]}'")
